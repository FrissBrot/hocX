"""Tests for the Word-Import importer recognizing a "dd.mm.yyyy - dd.mm.yyyy"-style
date range in a Termine-table cell (e.g. a holiday-plan row like "24.12.2025 -
06.01.2026  Weihnachtsferien") as one multi-day Termin instead of only ever reading
the range's start date and silently dropping the end date.

Covers three layers: the pure regex/parsing helpers (_extract_event_row for an
ordinary "events"-role table row, _extract_dates_with_counts for a Matrix "events"
cell), a real analyze() pass over a minimal .docx table, and a full commit() that
writes Event.event_end_date to the DB."""
from datetime import date
from io import BytesIO

from docx import Document

from app.models import Event
from app.schemas.word_import import WordImportCommit, WordImportEventCommit
from app.services import word_import_service as svc
from app.services.word_import_service import WordImportService
from tests.factories import make_event, make_tenant, make_template


# --- _extract_event_row (ordinary "events"-role table row) ------------------------


def test_extract_event_row_reads_a_hyphen_date_range():
    title, start, end = svc._extract_event_row(["24.12.2025 - 06.01.2026", "Weihnachtsferien"])
    assert title == "Weihnachtsferien"
    assert start == date(2025, 12, 24)
    assert end == date(2026, 1, 6)


def test_extract_event_row_reads_an_en_dash_date_range_in_one_cell():
    # Real documents (see export_service._matrix_event_row_value / the plain events
    # list export) round-trip a range using "–" (en dash), not a hyphen - both must
    # parse identically since re-importing a previously-exported protocol is a real
    # workflow.
    title, start, end = svc._extract_event_row(["3.", "14.05.2026 – 17.05.2026 Auffahrt"])
    assert title == "Auffahrt"
    assert start == date(2026, 5, 14)
    assert end == date(2026, 5, 17)


def test_extract_event_row_single_day_has_no_end_date():
    title, start, end = svc._extract_event_row(["08.12.2024", "Maria Empfängnis"])
    assert title == "Maria Empfängnis"
    assert start == date(2024, 12, 8)
    assert end is None


def test_extract_event_row_rejects_a_reversed_range_and_falls_back_to_the_start_date():
    # A malformed/reversed range (end before start) must not silently create an
    # invalid multi-day Event - event_service._build_event_entity would reject it
    # outright, so the importer degrades to a plain single-day row on the start date
    # instead of surfacing an opaque 500 at commit time.
    title, start, end = svc._extract_event_row(["06.01.2026 - 24.12.2025", "Kaputte Angabe"])
    assert title == "Kaputte Angabe"
    assert start == date(2026, 1, 6)
    assert end is None


# --- _extract_dates_with_counts (Matrix "events" cell) -----------------------------


def test_extract_dates_with_counts_reads_a_range_as_one_entry():
    results = svc._extract_dates_with_counts("24.12.2025 - 06.01.2026 (Weihnachtsferien)")
    assert results == [(date(2025, 12, 24), date(2026, 1, 6), None)]


def test_extract_dates_with_counts_keeps_the_participant_count_after_a_range():
    # Mirrors the plain single-date "18.10.2025 (7)" format
    # export_service._matrix_event_row_value writes - a range must support the same
    # trailing count annotation right after its end date.
    results = svc._extract_dates_with_counts("14.05.2026 – 17.05.2026 (12)")
    assert results == [(date(2026, 5, 14), date(2026, 5, 17), 12)]


def test_extract_dates_with_counts_does_not_merge_two_unrelated_standalone_dates():
    results = svc._extract_dates_with_counts("18.10.2025 (7) und 25.10.2025 (3)")
    assert results == [(date(2025, 10, 18), None, 7), (date(2025, 10, 25), None, 3)]


# --- analyze() over a real .docx table ---------------------------------------------


def _termine_docx(rows: list[list[str]]) -> bytes:
    document = Document()
    document.add_paragraph("Protokoll Hock vom 01.01.2026")
    # A real document always has its own Anwesenheit table - without one, analyze()'s
    # "no table was recognized as attendance at all -> assume table 0 is it" last-resort
    # fallback would hijack the Termine table below (table 0 here) back to "attendance"
    # (see analyze()'s comment above table_roles[parsed.tables[0].index] = "attendance").
    document.add_heading("Anwesenheit", level=1)
    attendance_table = document.add_table(rows=1, cols=2)
    attendance_table.rows[0].cells[0].text = "Name"
    attendance_table.rows[0].cells[1].text = "Status"
    document.add_heading("Termine", level=1)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Datum"
    table.rows[0].cells[1].text = "Anlass"
    for row_values in rows:
        row_cells = table.add_row().cells
        for cell, text in zip(row_cells, row_values):
            cell.text = text
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_analyze_surfaces_a_multiday_range_row_as_a_new_event_with_raw_end_date(db):
    tenant = make_tenant(db)
    template = make_template(db, tenant.id)
    template.protocol_number_pattern = "P-{n}"
    db.flush()

    docx_bytes = _termine_docx([["24.12.2025 - 06.01.2026", "Weihnachtsferien"]])
    analysis = WordImportService().analyze(
        db, tenant_id=tenant.id, template_id=template.id, protocol_date_hint=date(2026, 1, 1), raw_bytes=docx_bytes,
    )

    assert analysis.event_mappings, "expected the Termine row to be surfaced"
    mapping = analysis.event_mappings[0]
    assert mapping.raw_title == "Weihnachtsferien"
    assert mapping.raw_date == date(2025, 12, 24)
    assert mapping.raw_end_date == date(2026, 1, 6)
    assert mapping.status == "new"


def test_commit_creates_a_multiday_event_from_a_range_row(db):
    tenant = make_tenant(db)
    template = make_template(db, tenant.id)
    template.protocol_number_pattern = "P-{n}"
    db.flush()

    service = WordImportService()
    payload = WordImportCommit(
        template_id=template.id,
        protocol_date=date(2026, 1, 1),
        events=[
            WordImportEventCommit(
                approved=True,
                linked_event_id=None,
                final_title="Weihnachtsferien",
                final_date=date(2025, 12, 24),
                final_end_date=date(2026, 1, 6),
                raw_title="Weihnachtsferien",
                raw_date=date(2025, 12, 24),
                raw_end_date=date(2026, 1, 6),
            )
        ],
    )
    service.commit(db, tenant_id=tenant.id, user_id=1, payload=payload)

    created = db.query(Event).filter(Event.tenant_id == tenant.id, Event.title == "Weihnachtsferien").one()
    assert created.event_date == date(2025, 12, 24)
    assert created.event_end_date == date(2026, 1, 6)


def test_commit_update_writes_the_final_end_date_onto_an_existing_event(db):
    tenant = make_tenant(db)
    template = make_template(db, tenant.id)
    template.protocol_number_pattern = "P-{n}"
    db.flush()
    existing = make_event(db, tenant.id, title="Weihnachtsferien", event_date=date(2025, 12, 24))
    assert existing.event_end_date is None

    service = WordImportService()
    payload = WordImportCommit(
        template_id=template.id,
        protocol_date=date(2026, 1, 1),
        events=[
            WordImportEventCommit(
                approved=True,
                linked_event_id=existing.id,
                final_title="Weihnachtsferien",
                final_date=date(2025, 12, 24),
                final_end_date=date(2026, 1, 6),
                raw_title="Weihnachtsferien",
                raw_date=date(2025, 12, 24),
                raw_end_date=date(2026, 1, 6),
            )
        ],
    )
    service.commit(db, tenant_id=tenant.id, user_id=1, payload=payload)

    db.refresh(existing)
    assert existing.event_end_date == date(2026, 1, 6)
