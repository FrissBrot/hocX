"""Builds a realistic "Hock-Protokoll" test document (.docx and .pdf, byte-identical
content) for the word-import E2E tests in test_word_import_e2e.py. Both renderers are
driven off the exact same plain-data spec so the two formats can be asserted against
identical expectations - any divergence in results between the two is either a real
parser bug or a deliberate, documented format difference.

Layout mirrors the real hocX Hock-Protokoll shape confirmed against Timo's actual
documents (see project_hocx memory, "Word-Import-Tool" phases): a title/date line,
then a run of Heading-styled sections, some followed by a real two-column-or-wider
table (Anwesenheit/Termine/Ämtli/Matrix), some carrying only free "Label: Value" text
(Scharanlässe/Rückblick).
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import date
from io import BytesIO

from docx import Document


@dataclass
class TableSpec:
    heading: str
    header_cells: list[str]
    rows: list[list[str]]


@dataclass
class TextSpec:
    heading: str
    lines: list[str]


@dataclass
class ProtocolSpec:
    title: str
    attendance: TableSpec
    events: TableSpec
    list_table: TableSpec
    matrix: TableSpec
    form_text: TextSpec
    rueckblick_text: TextSpec
    extra_texts: list[TextSpec] = field(default_factory=list)


def default_spec(protocol_date: date = date(2026, 10, 18)) -> ProtocolSpec:
    d = protocol_date.strftime("%d.%m.%Y")
    return ProtocolSpec(
        title=f"Protokoll Hock vom {d}",
        attendance=TableSpec(
            heading="Anwesenheit",
            header_cells=["Name", "Status"],
            rows=[
                ["Timo Weber", ""],
                ["Nevio Muster", "entschuldigt"],
                ["Ganz Neue Person", ""],
            ],
        ),
        events=TableSpec(
            heading="Termine",
            header_cells=["Datum", "Anlass"],
            rows=[
                ["18.10.2026", "Herbsthock"],
                ["25.10.2026", "Vorstandssitzung"],
                ["01.11.2026", "Halloween-Party"],
            ],
        ),
        list_table=TableSpec(
            heading="Ämtli",
            header_cells=["Amt", "Person"],
            rows=[
                ["Feuer", "Timo Weber"],
                ["Fahrer", "Sandro Keller"],
            ],
        ),
        matrix=TableSpec(
            heading="Anwesenheitsmatrix",
            header_cells=["", "18.10.2026", "25.10.2026"],
            rows=[["Küchendienst", "Timo Weber", "Nevio Muster, Sandro Keller"]],
        ),
        form_text=TextSpec(
            heading="Scharanlässe",
            lines=[
                "Treffpunkt: Vor der Kirche",
                "Organisation: Timo Weber",
                "Wer geht: Timo Weber, Ganz Neue Person",
            ],
        ),
        rueckblick_text=TextSpec(
            heading="Rückblick Herbsthock",
            lines=["Positiv: Gutes Wetter, viele Teilnehmer.", "Negativ: Feuerholz war knapp."],
        ),
    )


def render_docx(spec: ProtocolSpec) -> bytes:
    document = Document()
    document.add_paragraph(spec.title)

    def _add_table(table_spec: TableSpec) -> None:
        if not table_spec.header_cells:
            return
        document.add_heading(table_spec.heading, level=1)
        table = document.add_table(rows=1, cols=len(table_spec.header_cells))
        for cell, text in zip(table.rows[0].cells, table_spec.header_cells):
            cell.text = text
        for row_values in table_spec.rows:
            row_cells = table.add_row().cells
            for cell, text in zip(row_cells, row_values):
                cell.text = text

    def _add_text(text_spec: TextSpec) -> None:
        if not text_spec.lines:
            return
        document.add_heading(text_spec.heading, level=1)
        for line in text_spec.lines:
            document.add_paragraph(line)

    _add_table(spec.attendance)
    _add_table(spec.events)
    _add_table(spec.list_table)
    _add_text(spec.form_text)
    _add_text(spec.rueckblick_text)
    _add_table(spec.matrix)
    for extra in spec.extra_texts:
        _add_text(extra)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_pdf(spec: ProtocolSpec) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    body_style = ParagraphStyle(name="Body", fontSize=11, leading=14)
    heading_style = ParagraphStyle(name="Heading", fontSize=16, leading=20, spaceBefore=10, spaceAfter=6)
    title_style = ParagraphStyle(name="Title", fontSize=11, leading=14)

    story = [Paragraph(spec.title, title_style), Spacer(1, 8)]

    def _table_flowable(table_spec: TableSpec) -> Table:
        data = [table_spec.header_cells] + table_spec.rows
        table = Table(data)
        table.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.75, colors.black),
                    ("FONTSIZE", (0, 0), (-1, -1), 11),
                ]
            )
        )
        return table

    def _add_table(table_spec: TableSpec) -> None:
        if not table_spec.header_cells:
            return
        story.append(Paragraph(table_spec.heading, heading_style))
        story.append(_table_flowable(table_spec))
        story.append(Spacer(1, 10))

    def _add_text(text_spec: TextSpec) -> None:
        if not text_spec.lines:
            return
        story.append(Paragraph(text_spec.heading, heading_style))
        for line in text_spec.lines:
            story.append(Paragraph(line, body_style))
        story.append(Spacer(1, 10))

    _add_table(spec.attendance)
    _add_table(spec.events)
    _add_table(spec.list_table)
    _add_text(spec.form_text)
    _add_text(spec.rueckblick_text)
    _add_table(spec.matrix)
    for extra in spec.extra_texts:
        _add_text(extra)

    buffer = BytesIO()
    SimpleDocTemplate(buffer, pagesize=A4).build(story)
    return buffer.getvalue()
