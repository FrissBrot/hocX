"""Word-import free-text sections must carry over enough formatting from the source
.docx that RichTextEditor (tiptap-markdown) renders them the way they looked in Word:
bold/italic runs, manual line breaks (Shift+Enter), and real paragraph breaks - see
ParsedSection.markdown_text / _paragraph_to_markdown in word_import_service.py.
"""
from io import BytesIO

from docx import Document
from docx.enum.text import WD_BREAK

from app.services.word_import_service import parse_docx


def _build_docx_with_formatted_section() -> bytes:
    document = Document()
    document.add_paragraph("Protokoll Hock vom 18.10.2026")
    document.add_heading("Rückblick", level=1)

    p1 = document.add_paragraph()
    p1.add_run("Das Fest war ")
    p1.add_run("sehr erfolgreich").bold = True
    p1.add_run(" und ")
    p1.add_run("kurzweilig").italic = True
    p1.add_run(".")

    p2 = document.add_paragraph()
    p2.add_run("Erste Zeile")
    p2.add_run().add_break(WD_BREAK.LINE)
    p2.add_run("Zweite Zeile (manueller Umbruch)")

    document.add_paragraph("Neuer Absatz nach der zweiten Zeile.")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_bold_and_italic_runs_become_markdown_markers():
    parsed = parse_docx(_build_docx_with_formatted_section())
    section = next(s for s in parsed.sections if s.heading == "Rückblick")

    assert "**sehr erfolgreich**" in section.markdown_text
    assert "*kurzweilig*" in section.markdown_text
    # Plain text used for classification/matching stays unformatted.
    assert "sehr erfolgreich" in section.text
    assert "**" not in section.text


def test_manual_line_break_becomes_hard_break_not_lost_or_merged():
    parsed = parse_docx(_build_docx_with_formatted_section())
    section = next(s for s in parsed.sections if s.heading == "Rückblick")

    assert "Erste Zeile\\\nZweite Zeile (manueller Umbruch)" in section.markdown_text


def test_paragraph_break_becomes_blank_line_not_single_newline():
    parsed = parse_docx(_build_docx_with_formatted_section())
    section = next(s for s in parsed.sections if s.heading == "Rückblick")

    assert "\n\nNeuer Absatz nach der zweiten Zeile." in section.markdown_text
    # A single "\n" between those two paragraphs would collapse into one run-on
    # paragraph when rendered by tiptap-markdown (breaks:false) instead of two.
    assert "\nNeuer Absatz" not in section.markdown_text.replace("\n\nNeuer Absatz", "")


def test_literal_markdown_characters_in_source_text_are_escaped():
    document = Document()
    document.add_paragraph("Protokoll Hock vom 18.10.2026")
    document.add_heading("Rückblick", level=1)
    document.add_paragraph("Preis: 5*3 Stück, Kontakt unter [siehe unten].")

    buffer = BytesIO()
    document.save(buffer)
    parsed = parse_docx(buffer.getvalue())
    section = next(s for s in parsed.sections if s.heading == "Rückblick")

    assert "5\\*3" in section.markdown_text
    assert "\\[siehe unten\\]" in section.markdown_text
