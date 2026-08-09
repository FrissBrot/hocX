from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass, field
from datetime import date
from difflib import SequenceMatcher
from io import BytesIO
from typing import Callable

import pdfplumber
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.cycle_utils import get_cycle_year
from app.repositories.participant_repository import participant_eligible_on
from app.services.isolated_parse import parse_document_isolated
from app.services.optimal_assignment import solve_optimal_assignment
from app.services.word_import_thresholds import adaptive_threshold
from app.models import (
    CycleConfig,
    ElementDefinition,
    ElementType,
    Event,
    ListDefinition,
    ListEntry,
    Participant,
    Protocol,
    ProtocolElement,
    ProtocolElementBlock,
    ProtocolText,
    Template,
    TemplateElement,
    TemplateParticipant,
    WordImportProfile,
    WordImportSuggestionOutcome,
)
from app.schemas.event import CycleAssignment, EventCreate, EventUpdate
from app.schemas.participant import ParticipantCreate
from app.schemas.protocol import ProtocolCreateFromTemplate, ProtocolUpdate
from app.schemas.word_import import (
    TablePreview,
    WordImportAnalysis,
    WordImportAttendanceCandidate,
    WordImportAttendanceMapping,
    WordImportCommit,
    WordImportEventCandidate,
    WordImportEventMapping,
    WordImportFormFieldValue,
    WordImportFormRow,
    WordImportListDefinitionOption,
    WordImportListEntryCandidate,
    WordImportListRowCommit,
    WordImportListRowMapping,
    WordImportMatrixCellMapping,
    WordImportMatrixColumnCandidate,
    WordImportMatrixOption,
    WordImportNameResolution,
    WordImportTextMapping,
    WordImportTextTarget,
)
from app.services.event_service import EventService
from app.services.participant_service import ParticipantService
from app.services.protocol_service import ProtocolService

_GERMAN_MONTHS = {
    "januar": 1, "februar": 2, "märz": 3, "maerz": 3, "april": 4, "mai": 5, "juni": 6,
    "juli": 7, "august": 8, "september": 9, "oktober": 10, "november": 11, "dezember": 12,
}
_EXCUSED_KEYWORDS = ("entschuldigt", "excused")
_ABSENT_KEYWORDS = ("abwesend", "unentschuldigt", "absent")
_LATE_KEYWORDS = ("verspätet", "verspaetet", "spät", "late")
_ATTENDANCE_TABLE_KEYWORDS = ("name", "teilnehmer", "mitglied", "anwesend")
_EVENT_TABLE_KEYWORDS = ("datum", "termin", "anlass", "sitzung")
_NAME_SPLIT_PATTERN = re.compile(r"[,/&]| und ")

_DATE_PATTERN = re.compile(r"(\d{1,2})\.\s?(\d{1,2})\.\s?(\d{2,4})")
_DATE_TEXT_PATTERN = re.compile(r"(\d{1,2})\.?\s+(" + "|".join(_GERMAN_MONTHS) + r")\s+(\d{4})", re.IGNORECASE)

_EVENT_MATCH_THRESHOLD = 0.8
_EVENT_CHANGE_THRESHOLD = 0.45
_LIST_NAME_MATCH_THRESHOLD = 0.5
_PARTICIPANT_MATCH_THRESHOLD = 0.6
_CANDIDATE_LIMIT = 5
_LIST_ENTRY_CANDIDATE_MIN_SCORE = 0.3
_MATRIX_ROW_MATCH_THRESHOLD = 0.5
_MATRIX_COLUMN_MATCH_THRESHOLD = 0.6
# Deliberately much stricter than the other thresholds above - a fuzzy signature match
# redirects an ENTIRE table's rows to a learned role/target, a much larger blast radius
# than a single name/event mismatch, so this only tolerates a near-identical header
# (e.g. one OCR/typo character, or one extra/missing normalized space) rather than a
# genuinely different table shape.
_TABLE_SIGNATURE_FUZZY_THRESHOLD = 0.92
# Score penalty applied to a candidate previously recorded as rejected for the exact
# same context (see WordImportProfile.mapping_config_json["rejected_candidates"]) - a
# demotion, not a hard exclusion, so a genuinely-still-correct match just needs a
# bigger margin over alternatives to win the auto-pick again.
_REJECTED_CANDIDATE_PENALTY = 0.15
# Deliberately excludes "&"/" und ": a raw List cell often holds several distinct
# values ("Felsenheim, Dolomiten Sport"), but "&"/" und " inside such a value usually
# joins a single entity's own name (e.g. company "Omlin & Partner Gmbh") rather than
# separating two different ones - splitting on them there corrupts the value. Order
# here also doubles as the tie-break preference in _select_list_row_variant when two
# delimiters score identically (comma before the riskier, much more eager "space").
_LIST_SPLIT_DELIMITERS: dict[str, re.Pattern[str]] = {
    "comma": re.compile(r","),
    "semicolon": re.compile(r";"),
    "slash": re.compile(r"/"),
    "newline": re.compile(r"\n"),
    "space": re.compile(r"\s+"),
}
_LIST_VARIANT_CONFIDENT_SCORE = 0.75
# Row types whose cell value can be resolved by this importer, reusing the exact same
# machinery as list columns/form fields (_match_names/_resolved_value_json). Matrix
# rows of type "event" or an embedded-block element_type_id (nested Terminliste/
# Anwesenheit/etc. inside a matrix cell) are recognized but skipped with a warning -
# out of scope for this first version, see plan.
_MATRIX_SUPPORTED_ROW_TYPES = {"text", "participant", "participants"}


def _best_fuzzy_signature_match(signature: str, profile_table_roles: dict[str, dict], threshold: float) -> dict | None:
    """Fallback for _resolve_table_role's profile lookup when no EXACT signature key
    matches - a template revision that adds/removes one column, or a minor OCR/typo
    difference in a header cell, would otherwise make a learned table-role mapping miss
    entirely and fall all the way back to the keyword heuristics. Only returns a match
    at or above `threshold` (see _TABLE_SIGNATURE_FUZZY_THRESHOLD's docstring for why
    that's set conservatively high)."""
    best_score = 0.0
    best_entry: dict | None = None
    for stored_signature, entry in profile_table_roles.items():
        score = _similarity(signature, stored_signature)
        if score > best_score:
            best_score = score
            best_entry = entry
    return best_entry if best_score >= threshold else None


@dataclass
class ParsedSection:
    heading: str
    text: str


@dataclass
class ParsedTable:
    index: int
    header_cells: list[str]
    rows: list[list[str]]
    preceding_heading: str | None
    known_role: str | None = None


@dataclass
class ParsedDocx:
    protocol_date: date | None
    title_hint: str | None
    sections: list[ParsedSection] = field(default_factory=list)
    tables: list[ParsedTable] = field(default_factory=list)


_UMLAUT_FOLD = str.maketrans({"ä": "a", "ö": "o", "ü": "u", "ß": "ss"})


def _fold_umlauts(text: str) -> str:
    return text.translate(_UMLAUT_FOLD)


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", _fold_umlauts(text.strip().lower()))


def _extract_date(text: str) -> date | None:
    candidate = _strip_leading_ordinal(text)
    match = _DATE_PATTERN.search(candidate)
    if match:
        day, month, year = match.groups()
        year_int = int(year) if len(year) == 4 else 2000 + int(year)
        try:
            return date(year_int, int(month), int(day))
        except ValueError:
            return None
    match = _DATE_TEXT_PATTERN.search(candidate.lower())
    if match:
        day, month_name, year = match.groups()
        try:
            return date(int(year), _GERMAN_MONTHS[month_name], int(day))
        except ValueError:
            return None
    return None


_PARTICIPANT_COUNT_PATTERN = re.compile(r"^\s*\((\d+)\)")


def _extract_dates_with_counts(text: str) -> list[tuple[date, int | None]]:
    """Finds every date mentioned anywhere in text, not just the first (unlike
    _extract_date) - used for a Matrix "events" cell, where multiple dates may be
    crammed onto one line/paragraph (e.g. Word soft line breaks, which python-docx
    does not reliably surface as "\\n" in Cell.text) instead of one date per line.
    Also captures a trailing "(N)" right after a date (e.g. "18.10.2025 (7)") as that
    date's participant count - the exact format export_service._matrix_event_row_value
    itself writes when event_show_participant_count is on, so an old exported/re-
    imported protocol round-trips its attendance counts. Deduplicates by date (first
    occurrence wins) while preserving first-seen order."""
    matches: list[tuple[int, int, date]] = []
    for match in _DATE_PATTERN.finditer(text):
        day, month, year = match.groups()
        year_int = int(year) if len(year) == 4 else 2000 + int(year)
        try:
            matches.append((match.start(), match.end(), date(year_int, int(month), int(day))))
        except ValueError:
            continue
    for match in _DATE_TEXT_PATTERN.finditer(text.lower()):
        day, month_name, year = match.groups()
        try:
            matches.append((match.start(), match.end(), date(int(year), _GERMAN_MONTHS[month_name], int(day))))
        except ValueError:
            continue
    matches.sort(key=lambda entry: entry[0])

    seen: set[date] = set()
    results: list[tuple[date, int | None]] = []
    for start_index, end_index, candidate in matches:
        if candidate in seen:
            continue
        seen.add(candidate)
        count_match = _PARTICIPANT_COUNT_PATTERN.match(text[end_index : end_index + 20])
        participant_count = int(count_match.group(1)) if count_match else None
        results.append((candidate, participant_count))
    return results


# Requires at least one space after the ordinal ("4. 25.09.2024" -> list item) so a
# tightly-written real date's day component ("13.09.2025", no space before the month)
# is never mistaken for a leading list number and mangled.
_LEADING_ORDINAL_PATTERN = re.compile(r"^\d{1,2}[.\)]\s+")


def _strip_leading_ordinal(text: str) -> str:
    return _LEADING_ORDINAL_PATTERN.sub("", text.strip(), count=1)


def _starts_with_date(text: str) -> bool:
    """Word templates sometimes reuse the same numbered Heading style for a
    dated sub-list inside a section (e.g. a holiday-calendar list numbered
    4-15 right in between real section headings "3." and "16."). A real
    section title in this document type never starts with a calendar date,
    so this overrides any style/bold signal that would otherwise misclassify
    those list rows as new sections."""
    remainder = _strip_leading_ordinal(text)
    return _DATE_PATTERN.match(remainder) is not None


def _is_heading(paragraph: DocxParagraph) -> bool:
    """Only a real Word heading style starts a new section - bold/large body text
    (e.g. an inline label like "Module:" within a section) must never split a
    section on its own, only genuine "Heading N"/"Überschrift N" paragraphs do."""
    text = paragraph.text.strip()
    if not text or len(text) > 120:
        return False
    if _starts_with_date(text):
        return False
    style_name = (paragraph.style.name if paragraph.style else "") or ""
    return style_name.lower().startswith(("heading", "überschrift", "ueberschrift", "titel"))


def _classify_status(marker: str) -> str | None:
    lowered = marker.strip().lower()
    if not lowered:
        return None
    if any(keyword in lowered for keyword in _EXCUSED_KEYWORDS):
        return "excused"
    if any(keyword in lowered for keyword in _ABSENT_KEYWORDS):
        return "absent"
    if any(keyword in lowered for keyword in _LATE_KEYWORDS):
        return "late"
    return None


def _token_jaccard(a: str, b: str) -> float:
    """Word-order-independent similarity component: Jaccard index over the two
    strings' whitespace-separated, umlaut-folded token sets (e.g. "Weber Timo" vs.
    "Timo Weber" -> 1.0, where SequenceMatcher's character-run comparison scores
    that pair poorly). Contributes nothing for single-token strings (Jaccard of two
    1-element sets is only ever 0.0 or 1.0, same as an exact/non-exact
    SequenceMatcher comparison there) - that case is handled separately by the
    nickname-equivalence table in _name_score."""
    tokens_a = set(_fold_umlauts(a.lower().strip()).split())
    tokens_b = set(_fold_umlauts(b.lower().strip()).split())
    if not tokens_a or not tokens_b:
        return 0.0
    return len(tokens_a & tokens_b) / len(tokens_a | tokens_b)


def _similarity(a: str, b: str) -> float:
    # max() with the token-Jaccard component is a per-pair monotonic improvement -
    # it can only raise a pair's own score, never lower it below what plain
    # SequenceMatcher already gave it. Note this does NOT guarantee two different
    # pairs' RELATIVE order is preserved (a pair whose Jaccard jumps more than
    # another's can flip which of the two now scores higher) - call sites that pick
    # a winner between competing pools by comparing their top scores (e.g.
    # _resolve_table_role's list-vs-matrix decision) should be covered by a
    # regression test rather than assumed unaffected.
    sequence_ratio = SequenceMatcher(None, _fold_umlauts(a.lower().strip()), _fold_umlauts(b.lower().strip())).ratio()
    return max(sequence_ratio, _token_jaccard(a, b))


def _iter_block_items(document):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, document)


def parse_docx(raw_bytes: bytes) -> ParsedDocx:
    """Heuristic extraction tailored to a single, consistent legacy Word template.

    Walks the document body in true reading order (not `document.paragraphs`/
    `document.tables` separately) so each table can be associated with whatever
    heading immediately precedes it, even when there's no paragraph text between
    them (e.g. a heading followed directly by a table, as in "2. Ämtli").
    """
    document = Document(BytesIO(raw_bytes))
    items = list(_iter_block_items(document))

    non_empty_paragraph_texts = [item.text.strip() for item in items if isinstance(item, DocxParagraph) and item.text.strip()]
    title_hint = non_empty_paragraph_texts[0] if non_empty_paragraph_texts else None
    protocol_date = None
    for text in non_empty_paragraph_texts[:15]:
        protocol_date = _extract_date(text)
        if protocol_date:
            break

    sections: list[ParsedSection] = []
    tables: list[ParsedTable] = []
    current_heading: str | None = None
    current_lines: list[str] = []
    table_index = 0
    for item in items:
        if isinstance(item, DocxParagraph):
            text = item.text.strip()
            if not text:
                continue
            if _is_heading(item):
                if current_heading is not None and current_lines:
                    sections.append(ParsedSection(heading=current_heading, text="\n".join(current_lines)))
                current_heading = text
                current_lines = []
            elif current_heading is not None:
                current_lines.append(text)
        else:
            raw_rows = [[cell.text.strip() for cell in row.cells] for row in item.rows]
            if raw_rows:
                header_cells, *data_rows = raw_rows
                tables.append(
                    ParsedTable(
                        index=table_index,
                        header_cells=header_cells,
                        rows=[row for row in data_rows if any(row)],
                        preceding_heading=current_heading,
                    )
                )
                table_index += 1
    if current_heading is not None and current_lines:
        sections.append(ParsedSection(heading=current_heading, text="\n".join(current_lines)))

    sections, tables = _reclassify_text_only_sections(sections, tables, table_index)
    return ParsedDocx(protocol_date=protocol_date, title_hint=title_hint, sections=sections, tables=tables)


def _reclassify_text_only_sections(
    sections: list[ParsedSection], tables: list[ParsedTable], table_index: int
) -> tuple[list[ParsedSection], list[ParsedTable]]:
    """Some legacy documents lay out what's structurally a Termine list or a
    role/person list as plain paragraph text instead of a real table (a Word table
    style, or - for a PDF - without any grid lines pdfplumber can detect) - this
    reclassifies such a section as an extra synthetic ParsedTable (see
    _classify_section_kind) instead of leaving it only importable as opaque free
    text. Shared by parse_docx and parse_pdf."""
    kept_sections: list[ParsedSection] = []
    for section in sections:
        lines = [line.strip() for line in section.text.split("\n") if line.strip()]
        kind = _classify_section_kind(lines)
        if kind == "text":
            kept_sections.append(section)
            continue
        if kind == "events":
            rows = [[line] for line in lines]
        else:
            rows = [list(pair) for pair in (_split_two_columns(line) for line in lines) if pair is not None]
        tables.append(
            ParsedTable(
                index=table_index,
                header_cells=[section.heading],
                rows=rows,
                preceding_heading=section.heading,
                known_role=kind,
            )
        )
        table_index += 1
    return kept_sections, tables


_PDF_LINE_CLUSTER_TOLERANCE = 3.0
_PDF_HEADING_SIZE_MARGIN = 0.5


def _cluster_pdf_lines(words: list[dict]) -> list[list[dict]]:
    """Groups words on one page into visual lines by vertical position. pdfplumber
    hands back individual words with float y-coordinates that can jitter by a
    fraction of a point even within the same printed line, so a plain
    round(top)-equality grouping would sometimes split one line into two."""
    ordered = sorted(words, key=lambda word: (word["top"], word["x0"]))
    lines: list[list[dict]] = []
    for word in ordered:
        if lines and abs(word["top"] - lines[-1][0]["top"]) <= _PDF_LINE_CLUSTER_TOLERANCE:
            lines[-1].append(word)
        else:
            lines.append([word])
    for line in lines:
        line.sort(key=lambda word: word["x0"])
    return lines


def parse_pdf(raw_bytes: bytes) -> ParsedDocx:
    """Heuristic extraction for text-based PDFs (e.g. a Word document exported/printed
    to PDF, the format hocX's word-import tool is designed for - not a scanned image,
    which has no text layer to extract at all). Mirrors parse_docx's approach, but a
    PDF has no paragraph style or real table object to rely on: headings are inferred
    from font size (a line printed noticeably larger than the document's body text,
    the same visual cue a human would use) and tables from pdfplumber's own
    line-based table detection, which works well for a bordered Word table exported
    to PDF (the normal case) but not for a borderless one - those fall through to the
    same plain-text list/table reclassification as an unrecognized docx table would."""
    page_items: list[list[tuple[float, str, object]]] = []
    body_size_votes: Counter = Counter()
    with pdfplumber.open(BytesIO(raw_bytes)) as pdf:
        for page in pdf.pages:
            tables = page.find_tables()
            table_bboxes = [table.bbox for table in tables]
            words = page.extract_words(extra_attrs=["size"], keep_blank_chars=False)
            for word in words:
                if word.get("size"):
                    body_size_votes[round(word["size"], 1)] += 1

            def _inside_table(word: dict, bboxes: list[tuple[float, float, float, float]] = table_bboxes) -> bool:
                return any(
                    bbox[0] - 3 <= word["x0"] and word["x1"] <= bbox[2] + 3
                    and bbox[1] - 3 <= word["top"] and word["bottom"] <= bbox[3] + 3
                    for bbox in bboxes
                )

            free_words = [word for word in words if not _inside_table(word)]
            items: list[tuple[float, str, object]] = []
            for line in _cluster_pdf_lines(free_words):
                text = " ".join(word["text"] for word in line).strip()
                if not text:
                    continue
                avg_size = sum(word.get("size", 0) for word in line) / len(line)
                items.append((line[0]["top"], "line", (text, avg_size)))
            for table in tables:
                rows = [[(cell or "").strip() for cell in row] for row in (table.extract() or [])]
                if rows:
                    items.append((table.bbox[1], "table", rows))
            items.sort(key=lambda entry: entry[0])
            page_items.append(items)

    body_size = body_size_votes.most_common(1)[0][0] if body_size_votes else 11
    heading_min_size = body_size + _PDF_HEADING_SIZE_MARGIN

    non_empty_texts = [payload[0] for items in page_items for _, kind, payload in items if kind == "line"]
    title_hint = non_empty_texts[0] if non_empty_texts else None
    protocol_date = None
    for text in non_empty_texts[:15]:
        protocol_date = _extract_date(text)
        if protocol_date:
            break

    sections: list[ParsedSection] = []
    tables: list[ParsedTable] = []
    current_heading: str | None = None
    current_lines: list[str] = []
    table_index = 0
    for items in page_items:
        for _, kind, payload in items:
            if kind == "line":
                text, avg_size = payload
                is_heading = len(text) <= 120 and not _starts_with_date(text) and avg_size >= heading_min_size
                if is_heading:
                    if current_heading is not None and current_lines:
                        sections.append(ParsedSection(heading=current_heading, text="\n".join(current_lines)))
                    current_heading = text
                    current_lines = []
                elif current_heading is not None:
                    current_lines.append(text)
            else:
                header_cells, *data_rows = payload
                tables.append(
                    ParsedTable(
                        index=table_index,
                        header_cells=header_cells,
                        rows=[row for row in data_rows if any(row)],
                        preceding_heading=current_heading,
                    )
                )
                table_index += 1
    if current_heading is not None and current_lines:
        sections.append(ParsedSection(heading=current_heading, text="\n".join(current_lines)))

    if not non_empty_texts and table_index == 0:
        raise ValueError("PDF enthält keinen erkennbaren Text (evtl. eingescannt) – Texterkennung wird nicht unterstützt.")

    sections, tables = _reclassify_text_only_sections(sections, tables, table_index)
    return ParsedDocx(protocol_date=protocol_date, title_hint=title_hint, sections=sections, tables=tables)


def parse_document(raw_bytes: bytes) -> ParsedDocx:
    """Dispatches to the right parser by sniffing the actual file signature - never
    trusts a filename/extension, consistent with FileService's own content-based mime
    check for stored word-import documents."""
    if raw_bytes[:5] == b"%PDF-":
        return parse_pdf(raw_bytes)
    return parse_docx(raw_bytes)


_ROLE_AMBIGUITY_EPSILON = 0.08


def _trial_resolve_list_confidence(
    table: ParsedTable, definition: ListDefinition, existing_entries: list[ListEntry], participants_by_id: dict[int, Participant]
) -> int:
    """Bounded probe for the list-vs-matrix tie-break below: how many of this table's
    rows would resolve confidently if it were actually a List against `definition` -
    reuses _select_list_row_variant's own scoring (_score_list_variant) rather than a
    separate metric, so "confident" means exactly what it already means everywhere
    else list rows are scored."""
    _, candidates, needs_manual, _ = _select_list_row_variant(table.rows, definition, existing_entries, participants_by_id, None)
    if needs_manual:
        return 0
    confident, _total = _score_list_variant(candidates, definition, existing_entries, participants_by_id)
    return confident


def _trial_resolve_matrix_confidence(table: ParsedTable, matrix: dict) -> int:
    """Bounded probe for the list-vs-matrix tie-break below: how many of this table's
    data rows have a row-label that clears _MATRIX_ROW_MATCH_THRESHOLD against one of
    `matrix`'s configured rows - the same row-resolution check analyze() itself
    performs per Matrix row, just counted here instead of applied."""
    matrix_rows = matrix["rows"]
    confident = 0
    for row_cells in table.rows:
        row_label_raw = row_cells[0] if row_cells else ""
        if not row_label_raw:
            continue
        best_score = max(
            (_similarity(row_label_raw, str(row.get("label") or row.get("title") or "")) for row in matrix_rows),
            default=0.0,
        )
        if best_score >= _MATRIX_ROW_MATCH_THRESHOLD:
            confident += 1
    return confident


def _resolve_table_role(
    table: ParsedTable,
    overrides: dict[int, dict],
    profile_table_roles: dict[str, dict],
    list_definitions: list[tuple[int, str]],
    matrices: list[dict],
    *,
    list_definitions_by_id: dict[int, ListDefinition] | None = None,
    list_entries_for: Callable[[int], list[ListEntry]] | None = None,
    participants_by_id: dict[int, Participant] | None = None,
) -> tuple[str, int | None, str | None, bool]:
    """Fourth return value is True when the role came from an explicit source (this
    call's manual override, or a learned profile signature match) rather than a
    heuristic guess - see the "first table defaults to attendance" fallback below,
    which must never clobber an explicit source just because it happens to sit at
    index 0. Third return value (matrix_key) is only ever set together with
    role == "matrix", mirroring how the second (list_definition_id) is only ever set
    together with role == "list"."""
    if table.index in overrides:
        entry = overrides[table.index]
        return entry.get("role", "ignore"), entry.get("list_definition_id"), entry.get("matrix_key"), True
    signature = _normalize(" | ".join(table.header_cells))
    if signature in profile_table_roles:
        entry = profile_table_roles[signature]
        return entry.get("role", "ignore"), entry.get("list_definition_id"), entry.get("matrix_key"), True
    fuzzy_entry = _best_fuzzy_signature_match(signature, profile_table_roles, _TABLE_SIGNATURE_FUZZY_THRESHOLD)
    if fuzzy_entry is not None:
        return fuzzy_entry.get("role", "ignore"), fuzzy_entry.get("list_definition_id"), fuzzy_entry.get("matrix_key"), True

    role = table.known_role
    if role is None:
        if any(keyword in signature for keyword in _ATTENDANCE_TABLE_KEYWORDS):
            role = "attendance"
        elif any(keyword in signature for keyword in _EVENT_TABLE_KEYWORDS):
            role = "events"

    # Matrix and List candidates are scored against the same heading together (not
    # list-first/matrix-first) so a table's preceding heading is matched against
    # whichever pool actually has the closer name, rather than an arbitrary priority
    # order between the two structurally different target kinds.
    if role is None and table.preceding_heading and (matrices or list_definitions):
        top_matrix_score = 0.0
        top_matrix: dict | None = None
        for matrix in matrices:
            score = _similarity(table.preceding_heading, matrix["title"])
            if score > top_matrix_score:
                top_matrix_score = score
                top_matrix = matrix
        top_list_score = 0.0
        top_list_id: int | None = None
        for list_id, name in list_definitions:
            score = _similarity(table.preceding_heading, name)
            if score > top_list_score:
                top_list_score = score
                top_list_id = list_id

        if top_matrix_score >= top_list_score:
            best_role, best_key, best_score = "matrix", (top_matrix["matrix_key"] if top_matrix else None), top_matrix_score
        else:
            best_role, best_key, best_score = "list", top_list_id, top_list_score

        # Ambiguous heading-only signal (both pools score within
        # _ROLE_AMBIGUITY_EPSILON of each other) - break the tie by which target
        # actually resolves more of this table's real rows, not by heading text
        # similarity alone (same "build variants, score, pick best" idea
        # _select_list_row_variant already applies to list-row grouping). Only runs
        # when the caller supplied the extra context needed to trial-resolve against
        # each candidate - callers that don't (or a table with no preceding_heading
        # match at all) keep today's pure heading-score behavior unchanged.
        if (
            top_matrix is not None and top_list_id is not None
            and abs(top_matrix_score - top_list_score) <= _ROLE_AMBIGUITY_EPSILON
            and list_definitions_by_id is not None and list_entries_for is not None and participants_by_id is not None
        ):
            definition = list_definitions_by_id.get(top_list_id)
            if definition is not None:
                list_confidence = _trial_resolve_list_confidence(
                    table, definition, list_entries_for(top_list_id), participants_by_id
                )
                matrix_confidence = _trial_resolve_matrix_confidence(table, top_matrix)
                if list_confidence > matrix_confidence:
                    best_role, best_key = "list", top_list_id
                elif matrix_confidence > list_confidence:
                    best_role, best_key = "matrix", top_matrix["matrix_key"]

        if best_role is not None and best_score >= _LIST_NAME_MATCH_THRESHOLD:
            if best_role == "matrix":
                return "matrix", None, best_key, False
            return "list", best_key, None, False

    if role == "list":
        # Synthetic list-shaped text section (see _classify_section_kind) - still try
        # to resolve WHICH List this belongs to via heading similarity before giving up
        # with no target. The block above only runs this same match while role is still
        # fully undetermined (role is None), so a table already known to be list-shaped
        # (its heading equals the section's own title, e.g. a heading "Ämtli" matched
        # exactly against an existing List named "Ämtli") previously fell straight
        # through to "no target found" without ever trying at all - a real bug, not
        # just a low-confidence miss. Scoped to list_definitions only (not matrices),
        # since the shape is already known to be list-like, not a cross-table.
        if table.preceding_heading and list_definitions:
            best_list_score = 0.0
            best_list_id: int | None = None
            for list_id, name in list_definitions:
                score = _similarity(table.preceding_heading, name)
                if score > best_list_score:
                    best_list_score = score
                    best_list_id = list_id
            if best_list_id is not None and best_list_score >= _LIST_NAME_MATCH_THRESHOLD:
                return "list", best_list_id, None, False
        return "list", None, None, False

    if role is not None:
        return role, None, None, False
    if len(table.header_cells) == 2:
        # A plausible two-column role/assignment table (like "Amt" / "Person") even
        # without a confident name match against an existing List - surfaced as
        # "list" with no target yet, rather than silently "ignore", so it's visible
        # in the review step and the user only has to pick which List it belongs to.
        return "list", None, None, False
    return "ignore", None, None, False


_DATE_RANGE_PATTERN = re.compile(
    r"\d{1,2}\.\s?\d{1,2}\.\s?\d{2,4}\s*[-–]\s*\d{1,2}\.\s?\d{1,2}\.\s?\d{2,4}"
)


def _extract_event_row(cells: list[str]) -> tuple[str, date | None]:
    for raw_cell in cells:
        # Strip a leading list ordinal ("14. ") first and consistently reuse that
        # stripped text for both date detection and remainder/title-building below -
        # otherwise the ordinal can fuse with the date itself ("4. 25.09.2024" being
        # misread as day=4/month=25) and derail everything downstream.
        cell = _strip_leading_ordinal(raw_cell)
        candidate_date = _extract_date(cell)
        if candidate_date is None:
            continue
        range_match = _DATE_RANGE_PATTERN.search(cell)
        if range_match:
            # "14.05.2026 – 17.05.2026 Auffahrt" - strip the whole range as one unit so
            # the end date doesn't linger in the title (a plain global sub would also
            # eat unrelated later dates in the same cell, which we don't want here).
            remainder = (cell[: range_match.start()] + cell[range_match.end() :]).strip(" -–:\t")
        else:
            remainder = _DATE_PATTERN.sub("", cell, count=1)
            remainder = _DATE_TEXT_PATTERN.sub("", remainder, count=1)
            remainder = remainder.strip(" -–:\t")
        if remainder:
            return remainder, candidate_date
        for other in cells:
            if other is raw_cell or not other.strip():
                continue
            if _extract_date(other) is None:
                return other.strip(), candidate_date
        return "", candidate_date
    for cell in cells:
        if cell.strip():
            return cell.strip(), None
    return "", None


def _split_two_columns(line: str) -> tuple[str, str] | None:
    if "\t" in line:
        parts = [part.strip() for part in line.split("\t") if part.strip()]
    else:
        parts = [part.strip() for part in re.split(r"\s{2,}", line) if part.strip()]
    if len(parts) == 2:
        return parts[0], parts[1]
    return None


def _classify_section_kind(lines: list[str]) -> str:
    """Some legacy templates lay out a Termine list or a role/person list as plain
    tab-aligned paragraph text instead of a real Word table - this lets such a
    section be treated like a table (offered as "events"/"list") instead of only
    ever being importable as opaque free text."""
    if len(lines) < 2:
        return "text"
    dated = sum(1 for line in lines if _starts_with_date(line))
    if dated / len(lines) >= 0.6:
        return "events"
    two_col = sum(1 for line in lines if _split_two_columns(line) is not None)
    if two_col / len(lines) >= 0.6:
        return "list"
    return "text"


_EVENT_REPEAT_MATCH_THRESHOLD = 0.5


def _strip_title_prefix(heading: str, element_title: str) -> str:
    """A "Rückblick"-style section's heading is normally "<Elementtitel> <Anlassname>"
    (e.g. "Rückblick Elternabend" for element title "Rückblick") - stripping the known
    element title isolates the Anlass name for the event-candidate search below. Falls
    back to the full heading if the title isn't actually a prefix of it (still a usable,
    if noisier, search string rather than no match at all)."""
    normalized_heading = _normalize(heading)
    normalized_title = _normalize(element_title)
    if normalized_title and normalized_heading.startswith(normalized_title):
        return heading[len(element_title):].strip(" :-–")
    return heading


def _score_event_candidate(title: str, raw_date: date | None, event: Event) -> float:
    title_score = _similarity(title, event.title)
    if raw_date is None:
        return 0.2 * title_score
    day_diff = abs((event.event_date - raw_date).days)
    if day_diff == 0:
        return 0.5 + 0.5 * title_score
    if day_diff <= 3:
        return 0.3 * max(0.0, 1 - day_diff / 3) + 0.4 * title_score
    return 0.2 * title_score


def _event_match_reason(title: str, raw_date: date | None, event: Event) -> str:
    """Short, human-readable justification for one WordImportEventCandidate, mirroring
    the same signals _score_event_candidate itself weighs - shown in the wizard so a
    reviewer can tell an exact-date-different-title match apart from a same-title-
    different-date one instead of only seeing an opaque number."""
    title_score = _similarity(title, event.title)
    if raw_date is not None and raw_date == event.event_date:
        return f"Datum exakt, Titel {round(title_score * 100)}% ähnlich"
    if raw_date is not None:
        day_diff = abs((event.event_date - raw_date).days)
        return f"Datum {day_diff}d abweichend, Titel {round(title_score * 100)}% ähnlich"
    return f"Titel {round(title_score * 100)}% ähnlich (kein Datum erkannt)"


def _text_match_reason(raw_text: str, candidate_label: str, score: float) -> str:
    """Generic reason string for candidates scored by plain _similarity (list entries,
    matrix columns, attendance names) - shared instead of duplicated per call site."""
    return f'"{raw_text}" zu "{candidate_label}": {round(score * 100)}% ähnlich'


# Maps a folded, lowercased nickname/short form to its canonical full first name, so
# _name_score can recognize e.g. "Sepp" and "Josef" as the same person's first token
# even though they share no substring/token overlap for _similarity to find. Deliberately
# a small, high-confidence starter set of common German/Swiss-German equivalences, not
# exhaustive - a wrong equivalence here would silently conflate two different real
# people, so only add more entries on confirmed real-document evidence.
_NICKNAME_TO_CANONICAL: dict[str, str] = {
    "sepp": "josef", "seppi": "josef", "josi": "josef",
    "hans": "johannes", "hansi": "johannes",
    "fritz": "friedrich", "fritzli": "friedrich",
    "heinz": "heinrich", "heiri": "heinrich",
    "res": "andreas", "resli": "andreas",
    "ueli": "ulrich", "uli": "ulrich",
    "hansueli": "hans-ulrich",
    "koebi": "jakob", "kobi": "jakob", "jockel": "jakob",
    "toni": "anton", "toeni": "anton",
    "gret": "margaretha", "gretli": "margaretha", "margrit": "margaretha",
    "trudi": "gertrud", "trudy": "gertrud",
    "vreni": "verena", "vreneli": "verena", "vroni": "verena",
    "hp": "hans-peter", "hanspeter": "hans-peter",
    "wisel": "alois", "wysel": "alois",
    "cheli": "karl", "chari": "karl", "carli": "karl",
    "lisi": "elisabeth", "lisbeth": "elisabeth", "betti": "elisabeth",
}


def _canonical_first_token(token: str) -> str:
    folded = _fold_umlauts(token.lower().strip())
    return _NICKNAME_TO_CANONICAL.get(folded, folded)


def _name_score(raw_name: str, display_name: str) -> float:
    # A raw name in a list/matrix cell is often just a first name (e.g. an informal
    # "Beisitzer: Nevio, Lino, Gian" column), while display_name is the full "Vorname
    # Nachname" - scored against the full name alone, SequenceMatcher penalizes the length
    # mismatch enough that a clean first-name match can fall under the threshold (e.g.
    # "Nevio" vs. "Nevio Kim Nguyen" scores ~0.45). Also score against just the display
    # name's first token and take the better of the two, so a first-name-only mention still
    # resolves as confidently as a full-name one.
    full_score = _similarity(raw_name, display_name)
    first_token = display_name.split(None, 1)[0] if display_name.strip() else display_name
    first_token_score = _similarity(raw_name, first_token)

    raw_parts = raw_name.split(None, 1)
    raw_first_token = raw_parts[0] if raw_parts else raw_name
    if _canonical_first_token(raw_first_token) != _canonical_first_token(first_token):
        return max(full_score, first_token_score)
    if len(raw_parts) <= 1:
        # raw_name is itself just a (nicknamed) bare first name - same "Nevio" vs.
        # "Nevio Kim Nguyen" bare-name case first_token_score exists for above, just
        # with a nickname spelling difference the plain character comparison wouldn't
        # otherwise resolve to a clean match.
        return max(full_score, first_token_score, 1.0)
    # raw_name carries its own surname/remainder too - the nickname only ever
    # substitutes for the FIRST token, so the surname must still independently match
    # display_name's remainder before this scores as a full match. Without this, "Sepp
    # Muster" would wrongly score a perfect match against "Josef Meier" purely off the
    # first-name nickname coincidence, ignoring the completely different surname.
    display_parts = display_name.split(None, 1)
    display_remainder = display_parts[1] if len(display_parts) > 1 else ""
    surname_score = _similarity(raw_parts[1], display_remainder) if display_remainder else 0.0
    return max(full_score, first_token_score, surname_score)


def _match_names(
    raw_text: str,
    participants: list[Participant],
    name_overrides: dict[str, int] | None = None,
    rejected_candidates: dict | None = None,
    match_threshold: float = _PARTICIPANT_MATCH_THRESHOLD,
) -> list[WordImportNameResolution]:
    overrides = name_overrides or {}
    # See A.4 / WordImportProfile.mapping_config_json["rejected_candidates"] - a
    # participant previously rejected for this exact raw name gets a score penalty
    # (not a hard exclusion) so a genuinely recurring wrong guess needs a bigger margin
    # over alternatives to win the auto-pick again, without ever disappearing from the
    # candidate list a human can still pick manually.
    rejected_by_name = rejected_candidates or {}
    names = [part.strip() for part in _NAME_SPLIT_PATTERN.split(raw_text) if part.strip()]
    resolutions: dict[int, int] = {}

    # Same two-phase strategy as the attendance matcher: saved overrides are claimed first in
    # document order, then the remaining names are resolved by GLOBAL best score across the
    # whole group (not left-to-right) so a mediocre match on an earlier name can't steal the
    # participant a later, better-matching name actually needs. Once a participant is claimed
    # by one name in this group, no other name in the same raw_text can also claim them (e.g.
    # "Nevio, Lino, Gian" must resolve to three distinct participants, not the same one three
    # times).
    used_participant_ids: set[int] = set()
    for index, name in enumerate(names):
        override_id = overrides.get(_normalize(name))
        if override_id is not None and override_id not in used_participant_ids:
            resolutions[index] = override_id
            used_participant_ids.add(override_id)

    def _score(index: int, participant: Participant) -> float:
        name = names[index]
        rejected_ids = set((rejected_by_name.get(f"name:{_normalize(name)}") or {}).get("rejected") or [])
        score = _name_score(name, participant.display_name)
        if participant.id in rejected_ids:
            score -= _REJECTED_CANDIDATE_PENALTY
        return score

    remaining_indices = [index for index in range(len(names)) if index not in resolutions]
    remaining_participants = [participant for participant in participants if participant.id not in used_participant_ids]
    # Globally optimal 1:1 assignment (Hungarian algorithm) over the remaining names x
    # remaining participants, replacing the previous "collect all pairs >= 0.4, sort
    # globally, greedily claim >= threshold" approximation - see optimal_assignment.py
    # docstring for why a greedy pass can strand a name with a worse match than the
    # optimal solution finds. The old 0.4 pre-filter is subsumed: the solver already
    # considers the whole matrix at once, so a separate collection threshold below the
    # real acceptance threshold serves no purpose here.
    for assignment in solve_optimal_assignment(
        remaining_indices, remaining_participants, _score, min_score=match_threshold
    ):
        resolutions[assignment.row] = assignment.col.id
        used_participant_ids.add(assignment.col.id)

    return [
        WordImportNameResolution(
            raw_name=name, participant_id=resolutions.get(index),
            # Set to the SAME value as participant_id at construction time - this is the
            # one and only place a WordImportNameResolution's suggestion is ever set; the
            # wizard's edit handlers only touch participant_id afterward (object-spread
            # preserves this field), so a divergence at commit() time means the human
            # picked something other than what analyze() originally suggested here.
            originally_suggested_participant_id=resolutions.get(index),
            # Same top-3-above-0.4 shape as the attendance table's own candidates (see
            # below) - even a name that never cleared match_threshold anywhere still
            # carries its best near-misses, so the wizard's recurring-name clarifier has
            # something to suggest instead of an empty search box.
            candidates=[
                WordImportAttendanceCandidate(
                    participant_id=participant.id, score=round(score, 3),
                    reason=_text_match_reason(name, participant.display_name, score),
                )
                for score, participant in sorted(
                    ((_score(index, participant), participant) for participant in participants),
                    key=lambda entry: entry[0],
                    reverse=True,
                )[:3]
                if score >= 0.4
            ],
        )
        for index, name in enumerate(names)
    ]


def _build_column_value(
    value_type: str,
    raw_text: str,
    participants: list[Participant],
    name_overrides: dict[str, int] | None = None,
    rejected_candidates: dict | None = None,
    match_threshold: float = _PARTICIPANT_MATCH_THRESHOLD,
) -> tuple[dict, list[WordImportNameResolution]]:
    if value_type == "text":
        text_value = raw_text.strip()
        return ({"text_value": text_value} if text_value else {}), []
    if value_type in ("participant", "participants"):
        resolutions = _match_names(raw_text, participants, name_overrides, rejected_candidates, match_threshold)
        matched_ids = [resolution.participant_id for resolution in resolutions if resolution.participant_id is not None]
        if value_type == "participant":
            return ({"participant_id": matched_ids[0]} if matched_ids else {}), resolutions
        return ({"participant_ids": matched_ids} if matched_ids else {}), resolutions
    return {}, []


_FORM_FIELD_LINE_PATTERN = re.compile(r"^(.+?):\s*(.*)$")
_FORM_FIELD_MATCH_THRESHOLD = 0.5


def _parse_form_fields(text: str, rows: list[dict]) -> dict[str, str]:
    """Extracts "Label: Value" lines from a section's raw text and fuzzy-matches each
    label against a form block's configured row labels (e.g. "Treffpunkt: Vor der
    Kirche" -> the "Treffpunkt" row) - the same "Label:" convention this app's own
    event-repeat text blocks already use by default (e.g. "Positiv:"/"Negativ:" in a
    Rückblick block), just split across dedicated form rows instead of one text blob.
    Only ':' is treated as a separator (not '-'/'–') to avoid misreading ordinary prose
    containing a dash as a field line. First matching line wins per row."""
    values: dict[str, str] = {}
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        match = _FORM_FIELD_LINE_PATTERN.match(line)
        if not match:
            continue
        raw_label, raw_value = match.groups()
        best_row_id: str | None = None
        best_score = 0.0
        for row in rows:
            score = _similarity(raw_label, str(row.get("label") or ""))
            if score > best_score:
                best_score = score
                best_row_id = str(row.get("id"))
        if best_row_id is not None and best_score >= _FORM_FIELD_MATCH_THRESHOLD and best_row_id not in values:
            values[best_row_id] = raw_value.strip()
    return values


def _resolved_value_json(value_type: str, raw_text: str, names: list) -> dict:
    """Turns an already name-resolved value (raw_text + a list of objects exposing
    .raw_name/.participant_id, e.g. WordImportNameResolution/WordImportFormFieldValue
    entries) into the {"text_value"|"participant_id"|"participant_ids": ...} shape stored
    on a list entry or form-block row - shared by the list-row commit path and the
    form-block text commit path below, which both resolve names the same way in the
    wizard before commit() runs."""
    if value_type in ("participant", "participants"):
        ids = [name.participant_id for name in names if name.participant_id is not None]
        if value_type == "participant":
            return {"participant_id": ids[0]} if ids else {}
        return {"participant_ids": ids} if ids else {}
    if value_type == "text":
        text_value = raw_text.strip()
        return {"text_value": text_value} if text_value else {}
    return {}


def _display_value(value_type: str, value_json: dict, participants_by_id: dict[int, Participant]) -> str:
    if value_type == "text":
        return str(value_json.get("text_value", ""))
    if value_type == "participant":
        participant_id = value_json.get("participant_id")
        participant = participants_by_id.get(participant_id) if participant_id is not None else None
        return participant.display_name if participant else ""
    if value_type == "participants":
        ids = value_json.get("participant_ids") or []
        return ", ".join(participants_by_id[pid].display_name for pid in ids if pid in participants_by_id)
    return ""


def _nearest_previous_protocol_id(db: Session, *, tenant_id: int, template_id: int, protocol_date: date) -> int | None:
    return db.scalar(
        select(Protocol.id)
        .where(
            Protocol.tenant_id == tenant_id,
            Protocol.template_id == template_id,
            Protocol.protocol_date < protocol_date,
        )
        .order_by(Protocol.protocol_date.desc(), Protocol.id.desc())
        .limit(1)
    )


def _list_snapshot_entries_by_id(db: Session, *, protocol_id: int, list_definition_id: int) -> dict[int, dict]:
    """entry_id -> frozen {"column_one_value", "column_two_value"} from a previous
    protocol's whole-list "form" block for this list, or {} if it has none. Used as the
    diff baseline for matched/changed instead of today's live value, since an old
    document should be compared against what was true back then, not today's data."""
    configs = db.execute(
        select(ProtocolElementBlock.configuration_snapshot_json)
        .join(ProtocolElement, ProtocolElement.id == ProtocolElementBlock.protocol_element_id)
        .where(ProtocolElement.protocol_id == protocol_id)
    ).scalars()
    for config in configs:
        if not config or config.get("linked_list_id") != list_definition_id:
            continue
        entries = (config.get("list_snapshot") or {}).get("entries") or []
        return {entry["id"]: entry for entry in entries if isinstance(entry, dict) and entry.get("id") is not None}
    return {}


def _template_linked_list_ids(db: Session, *, template_id: int, form_type_id: int | None) -> set[int]:
    """List ids the template already has a whole-list "form" block for - the only lists
    a word-import can write a snapshot into (see WordImportListRowMapping)."""
    if form_type_id is None:
        return set()
    ids: set[int] = set()
    definitions = db.execute(
        select(ElementDefinition)
        .join(TemplateElement, TemplateElement.element_definition_id == ElementDefinition.id)
        .where(TemplateElement.template_id == template_id)
    ).scalars()
    for definition in definitions:
        for block in (definition.configuration_json or {}).get("blocks", []):
            if block.get("element_type_id") != form_type_id:
                continue
            linked_list_id = (block.get("configuration_json") or {}).get("linked_list_id")
            if linked_list_id:
                ids.add(int(linked_list_id))
    return ids


def _template_matrices(db: Session, *, template_id: int, matrix_type_id: int | None) -> list[dict]:
    """Matrix block configs (rows/columns/mode/auto_source) available as import targets
    in this template. Keyed the same way as text_targets/block_by_key -
    "{template_element_id}:{block.sort_index}" - so a resolved match can be looked up
    against the live protocol's blocks in WordImportService.commit without a separate
    identity scheme (mirrors _template_linked_list_ids's join pattern)."""
    if matrix_type_id is None:
        return []
    matrices: list[dict] = []
    definitions = db.execute(
        select(TemplateElement, ElementDefinition)
        .join(ElementDefinition, ElementDefinition.id == TemplateElement.element_definition_id)
        .where(TemplateElement.template_id == template_id)
        .order_by(TemplateElement.sort_index.asc())
    ).all()
    for template_element, definition in definitions:
        for block in (definition.configuration_json or {}).get("blocks", []):
            if block.get("element_type_id") != matrix_type_id:
                continue
            sort_index = block.get("sort_index")
            config = block.get("configuration_json") or {}
            matrices.append(
                {
                    "matrix_key": f"{template_element.id}:{sort_index}",
                    "template_element_id": template_element.id,
                    "sort_index": sort_index,
                    "title": definition.title,
                    "mode": config.get("mode") or "manual",
                    "auto_source": config.get("auto_source") or {},
                    "rows": config.get("rows") or config.get("field_rows") or [],
                    "columns": config.get("columns") or config.get("matrix_columns") or [],
                }
            )
    return matrices


def _value_key(value_type: str, value_json: dict) -> str:
    if value_type == "text":
        return _normalize(str(value_json.get("text_value", "")))
    if value_type == "participant":
        return str(value_json.get("participant_id", ""))
    if value_type == "participants":
        return ",".join(str(pid) for pid in sorted(value_json.get("participant_ids", []) or []))
    if value_type == "event":
        return str(value_json.get("event_id", ""))
    return ""


@dataclass
class ListRowCandidate:
    """One resolved (column_one_raw, column_two_raw) pair produced by a grouping
    variant - source_row_index points back at the originating document row (several
    candidates from an "explode" variant can share the same source_row_index).
    group_filled marks a value that was inferred (fill-down/exploded repeat) rather
    than literally present in that document cell, so the wizard can flag it for
    review."""

    source_row_index: int
    column_one_raw: str
    column_two_raw: str
    group_filled: bool = False


def _flat_list_rows(rows: list[list[str]]) -> list[ListRowCandidate]:
    """The importer's original behaviour: row N's cell 0/1 map straight to
    column_one_raw/column_two_raw, rows with an empty first cell are dropped. Kept as
    its own variant (rather than inlined in the loop) so it can be scored against the
    alternatives below and stays the default when nothing scores better."""
    candidates: list[ListRowCandidate] = []
    for row_index, cells in enumerate(rows):
        column_one_raw = cells[0] if len(cells) > 0 else ""
        column_two_raw = cells[1] if len(cells) > 1 else ""
        if not column_one_raw:
            continue
        candidates.append(ListRowCandidate(row_index, column_one_raw, column_two_raw))
    return candidates


def _swapped_list_rows(rows: list[list[str]]) -> list[ListRowCandidate]:
    """Same as _flat_list_rows with the two cells swapped - covers documents whose
    table simply orders the two values opposite to the target list's own column
    order."""
    candidates: list[ListRowCandidate] = []
    for row_index, cells in enumerate(rows):
        column_one_raw = cells[1] if len(cells) > 1 else ""
        column_two_raw = cells[0] if len(cells) > 0 else ""
        if not column_one_raw:
            continue
        candidates.append(ListRowCandidate(row_index, column_one_raw, column_two_raw))
    return candidates


def _fill_down_list_rows(rows: list[list[str]]) -> list[ListRowCandidate]:
    """Handles Word tables where the first column is only filled on a group's first
    row (vertically merged/rowspan-styled in Word) and left blank on the group's other
    rows - python-docx surfaces those continuation cells as empty text, which
    _flat_list_rows silently drops today. Carries the previous non-empty first-cell
    value down onto blank-first-cell rows instead."""
    candidates: list[ListRowCandidate] = []
    last_column_one = ""
    for row_index, cells in enumerate(rows):
        column_one_raw = (cells[0] if len(cells) > 0 else "").strip()
        column_two_raw = cells[1] if len(cells) > 1 else ""
        group_filled = False
        if column_one_raw:
            last_column_one = column_one_raw
        elif last_column_one and column_two_raw.strip():
            column_one_raw = last_column_one
            group_filled = True
        if not column_one_raw:
            continue
        candidates.append(ListRowCandidate(row_index, column_one_raw, column_two_raw, group_filled))
    return candidates


def _present_delimiters(rows: list[list[str]], raw_index: int) -> list[str]:
    """Which of _LIST_SPLIT_DELIMITERS actually occur in this document column, i.e.
    are "smart"ly worth trying as an explode split - a delimiter only counts if it
    splits at least one row's cell into 2+ non-empty parts, so e.g. "slash" is never
    tried as a variant when no cell in this column contains a "/" at all. Returned in
    _LIST_SPLIT_DELIMITERS's declared order (also the tie-break preference order)."""
    present: list[str] = []
    for name, pattern in _LIST_SPLIT_DELIMITERS.items():
        for cells in rows:
            cell = cells[raw_index] if len(cells) > raw_index else ""
            parts = [part.strip() for part in pattern.split(cell) if part.strip()]
            if len(parts) >= 2:
                present.append(name)
                break
    return present


def _exploded_list_rows(
    rows: list[list[str]], *, text_raw_index: int, text_target: str, delimiter: re.Pattern[str]
) -> list[ListRowCandidate]:
    """Handles documents that group by one column and cram every value belonging to
    that group into a single, multi-value cell of the other column - e.g. one document
    row "Enea | Omlin & Partner Gmbh, Felsenheim, Dolomiten Sport" really means three
    target-list rows, each pairing one sponsor with "Enea". text_raw_index (0 or 1) is
    which raw document cell holds the multi-value text to split on `delimiter` (one of
    _LIST_SPLIT_DELIMITERS, chosen by the caller per _present_delimiters);
    text_target ("one"/"two") is which target list column those split values belong
    in - the other raw cell is the group value, repeated onto every row produced from
    that source row. Only ever called when exactly one of the list's two columns is
    value_type "text" (see _build_list_row_variants), since splitting into several
    independent values only makes sense for freeform text."""
    group_raw_index = 1 - text_raw_index
    candidates: list[ListRowCandidate] = []
    for row_index, cells in enumerate(rows):
        text_cell = cells[text_raw_index] if len(cells) > text_raw_index else ""
        group_value = (cells[group_raw_index] if len(cells) > group_raw_index else "").strip()
        if not group_value or not text_cell.strip():
            continue
        parts = [part.strip() for part in delimiter.split(text_cell) if part.strip()]
        if len(parts) < 2:
            continue
        for part in parts:
            if text_target == "one":
                candidates.append(ListRowCandidate(row_index, part, group_value, group_filled=True))
            else:
                candidates.append(ListRowCandidate(row_index, group_value, part, group_filled=True))
    return candidates


def _build_list_row_variants(rows: list[list[str]], definition: ListDefinition) -> dict[str, list[ListRowCandidate]]:
    """Builds every plausible interpretation of a document table's rows against one
    target ListDefinition - the caller (analyze()) scores each against the list's live
    entries and picks whichever produces the most confident matches ("meiste
    Zuweisungen"). Explode variants are generated per delimiter actually present in
    the data (see _present_delimiters) rather than a single fixed split pattern, so a
    document using "/" or plain whitespace to separate values gets exactly as fair a
    shot as one using commas - keyed "explode:<delimiter>"/"explode_swap:<delimiter>"
    so the caller can tell callers/the wizard exactly which split won."""
    variants: dict[str, list[ListRowCandidate]] = {
        "flat": _flat_list_rows(rows),
        "fill_down": _fill_down_list_rows(rows),
        "swap": _swapped_list_rows(rows),
    }
    text_target: str | None = None
    if definition.column_one_value_type == "text" and definition.column_two_value_type != "text":
        text_target = "one"
    elif definition.column_two_value_type == "text" and definition.column_one_value_type != "text":
        text_target = "two"
    if text_target is not None:
        for text_raw_index, prefix in ((0, "explode"), (1, "explode_swap")):
            for delimiter_name in _present_delimiters(rows, text_raw_index):
                exploded = _exploded_list_rows(
                    rows,
                    text_raw_index=text_raw_index,
                    text_target=text_target,
                    delimiter=_LIST_SPLIT_DELIMITERS[delimiter_name],
                )
                if exploded:
                    variants[f"{prefix}:{delimiter_name}"] = exploded
    return variants


def _score_list_variant(
    candidates: list[ListRowCandidate],
    definition: ListDefinition,
    existing_entries: list[ListEntry],
    participants_by_id: dict[int, Participant],
) -> tuple[int, int]:
    """Scores one variant's candidate rows against the list's live entries: how many
    resolve confidently, i.e. the best _similarity() against any existing entry's
    column_one display value clears _LIST_VARIANT_CONFIDENT_SCORE. This is the "meiste
    Zuweisungen" signal used to pick the best grouping interpretation of an ambiguous
    document table - reuses the exact same _similarity/_display_value comparison the
    main per-row candidate list below already does, just condensed to a best-score
    count instead of a full sorted candidate list."""
    confident = 0
    for candidate in candidates:
        best_score = 0.0
        for entry in existing_entries:
            score = _similarity(
                candidate.column_one_raw,
                _display_value(definition.column_one_value_type, entry.column_one_value_json or {}, participants_by_id),
            )
            if score > best_score:
                best_score = score
        if best_score >= _LIST_VARIANT_CONFIDENT_SCORE:
            confident += 1
    return confident, len(candidates)


# Tie-break preference when two variants score identically confident matches - flat/
# swap/fill_down (no splitting at all) come first, then whichever _LIST_SPLIT_DELIMITERS
# order the exploded variant's delimiter has (comma before the much more eager "space").
_VARIANT_BASE_PREFERENCE = ["flat", "swap", "fill_down"]


def _variant_preference_rank(name: str) -> int:
    if name in _VARIANT_BASE_PREFERENCE:
        return _VARIANT_BASE_PREFERENCE.index(name)
    delimiter_name = name.rsplit(":", 1)[-1] if ":" in name else None
    delimiter_order = list(_LIST_SPLIT_DELIMITERS.keys())
    offset = delimiter_order.index(delimiter_name) if delimiter_name in delimiter_order else len(delimiter_order)
    return len(_VARIANT_BASE_PREFERENCE) + offset


def _select_list_row_variant(
    rows: list[list[str]],
    definition: ListDefinition,
    existing_entries: list[ListEntry],
    participants_by_id: dict[int, Participant],
    forced_strategy: str | None,
) -> tuple[str, list[ListRowCandidate], bool, list[str]]:
    """Returns (strategy_name, chosen_rows, needs_manual_grouping, available_strategies).
    needs_manual_grouping is True exactly when the list has no live entries yet to
    score variants against - in that case automatic selection is meaningless and the
    wizard must let the user pick a strategy manually (see word_import.py
    TablePreview.needs_manual_grouping); available_strategies lists every variant name
    _build_list_row_variants actually produced for this table's data, so the wizard's
    manual picker only ever offers choices that exist for real."""
    variants = _build_list_row_variants(rows, definition)
    available_strategies = list(variants.keys())
    if forced_strategy and forced_strategy in variants:
        return forced_strategy, variants[forced_strategy], False, available_strategies
    if not existing_entries:
        return "flat", variants["flat"], True, available_strategies
    scored = {
        name: _score_list_variant(candidates, definition, existing_entries, participants_by_id)
        for name, candidates in variants.items()
    }

    def _rank(name: str) -> tuple[int, int, int]:
        confident, total = scored[name]
        return (confident, -_variant_preference_rank(name), total)

    best_name = max(scored, key=_rank)
    return best_name, variants[best_name], False, available_strategies


def _positional_matrix_column_resolution(
    doc_column_labels: list[str], targets: list[dict]
) -> dict[int, tuple[str | None, list[WordImportMatrixColumnCandidate]]]:
    """Alternative to label-similarity matching for a Matrix's fixed, template-
    configured columns: document column N maps straight to `targets[N]` in configured
    order, no text comparison at all - for documents whose column headers are missing
    or garbled (e.g. a borderless table pdfplumber only partially recovered) but whose
    column COUNT and ORDER still matches the template. Only ever tried as a guarded
    fallback (see call site) when label-matching already resolved fewer than half of
    the table's labeled columns - a reordered-columns document would make this a wrong
    guess, so it must never override an already-adequate label match."""
    resolution: dict[int, tuple[str | None, list[WordImportMatrixColumnCandidate]]] = {}
    for col_idx, label in enumerate(doc_column_labels):
        if not label or col_idx >= len(targets):
            continue
        target = targets[col_idx]
        column_key = str(target.get("id"))
        column_label = str(target.get("title") or "")
        resolution[col_idx] = (
            column_key,
            [
                WordImportMatrixColumnCandidate(
                    column_key=column_key, label=column_label, score=1.0, reason="Positionale Zuordnung (Spaltenreihenfolge)"
                )
            ],
        )
    return resolution


def _count_confident_matrix_columns(resolution: dict[int, tuple[str | None, list]]) -> int:
    return sum(1 for column_key, _candidates in resolution.values() if column_key is not None)


class WordImportService:
    def analyze(
        self,
        db: Session,
        *,
        tenant_id: int,
        template_id: int,
        protocol_date_hint: date | None,
        raw_bytes: bytes,
        table_role_overrides: dict[int, dict] | None = None,
        in_memory_profile_hints: dict | None = None,
    ) -> WordImportAnalysis:
        # Isoliert statt parse_document(raw_bytes) direkt - siehe isolated_parse.py: ein
        # pathologisches/bösartig komprimiertes Dokument wird nach Timeout hart
        # abgebrochen statt den aufrufenden Worker unbegrenzt zu blockieren.
        parsed = parse_document_isolated(raw_bytes)
        protocol_date = protocol_date_hint or parsed.protocol_date
        warnings: list[str] = []

        profile = db.execute(
            select(WordImportProfile).where(
                WordImportProfile.tenant_id == tenant_id, WordImportProfile.template_id == template_id
            )
        ).scalar_one_or_none()
        # See C.11 / WordImportQueueService's batch-consensus pass - in_memory_profile_hints
        # (same shape as mapping_config_json) is merged on top of the persisted profile for
        # THIS analyze() call only, never written back to WordImportProfile itself. Lets a
        # same-batch upload's confidently-resolved decisions help sibling documents in that
        # batch that failed to resolve the same raw text on their own, without waiting for
        # any of them to actually be committed first (the DB profile only ever learns from
        # real commits).
        profile_config = dict(profile.mapping_config_json or {}) if profile else {}
        if in_memory_profile_hints:
            profile_config = {
                "heading_to_target": {**profile_config.get("heading_to_target", {}), **in_memory_profile_hints.get("heading_to_target", {})},
                "table_roles_by_signature": {
                    **profile_config.get("table_roles_by_signature", {}), **in_memory_profile_hints.get("table_roles_by_signature", {})
                },
                "participant_name_overrides": {
                    **profile_config.get("participant_name_overrides", {}), **in_memory_profile_hints.get("participant_name_overrides", {})
                },
                "rejected_candidates": profile_config.get("rejected_candidates", {}),
            }
        heading_to_target = profile_config.get("heading_to_target", {})
        table_roles_by_signature = profile_config.get("table_roles_by_signature", {})
        participant_name_overrides = profile_config.get("participant_name_overrides", {})
        # See A.4 / WordImportService.commit's _log_outcome - {"{prefix}:{context}":
        # {"rejected": [ids], "chosen": id}}, consulted below (_rejected_ids_for) to
        # demote a candidate that was already wrong for this exact context once before,
        # rather than silently repeating the same mistake every time it recurs.
        rejected_candidates = profile_config.get("rejected_candidates", {})
        profile_applied = bool(heading_to_target or table_roles_by_signature or participant_name_overrides)

        def _rejected_ids_for(rejection_key: str) -> set:
            return set((rejected_candidates.get(rejection_key) or {}).get("rejected") or [])

        # See B.7 / word_import_thresholds.adaptive_threshold - resolved ONCE per
        # analyze() call (not read from inside the otherwise DB-free scoring helpers
        # further below, which stay pure/independently unit-testable) and threaded down
        # as an explicit parameter at each call site, so the source of a threshold value
        # is always traceable rather than a hidden DB read behind a "pure" function.
        # "participant_match" (the attendance-roster signal) is used as the shared proxy
        # for _PARTICIPANT_MATCH_THRESHOLD everywhere a name is matched (attendance,
        # matrix/list/form names all resolve through the same underlying scoring) - a
        # separate per-signal-type threshold for each name-matching context is not worth
        # the added complexity without concrete evidence they should diverge.
        participant_match_threshold = adaptive_threshold(
            db, tenant_id=tenant_id, template_id=template_id, signal_type="participant_match", default=_PARTICIPANT_MATCH_THRESHOLD
        )
        # Feeds B.5's event Hungarian assignment min_score - the looser "is this row
        # linked to an existing event AT ALL" gate, not the stricter matched-vs-changed
        # distinction (_EVENT_MATCH_THRESHOLD stays a fixed constant, see B.7 plan).
        event_change_threshold = adaptive_threshold(
            db, tenant_id=tenant_id, template_id=template_id, signal_type="event_match", default=_EVENT_CHANGE_THRESHOLD
        )

        list_definition_rows = list(
            db.execute(
                select(ListDefinition).where(ListDefinition.tenant_id == tenant_id, ListDefinition.is_active.is_(True))
            ).scalars()
        )
        list_definitions_for_matching = [(item.id, item.name) for item in list_definition_rows]
        list_definitions_by_id = {item.id: item for item in list_definition_rows}

        matrix_type_id = db.scalar(select(ElementType.id).where(ElementType.code == "matrix"))
        matrices_for_matching = _template_matrices(db, template_id=template_id, matrix_type_id=matrix_type_id)
        matrices_by_key = {matrix["matrix_key"]: matrix for matrix in matrices_for_matching}

        # Hoisted ahead of the table-role resolution loop below (originally fetched much
        # later in this method) so _resolve_table_role's optional list-vs-matrix
        # ambiguity tie-break (see B.6) has real data to trial-resolve against instead of
        # only ever falling back to pure heading-score comparison. list_entries_by_list_id
        # is the same lazy per-list-id cache the Matrix "list" auto-source section further
        # below already needs - declared once here and shared, not rebuilt twice.
        participants = list(
            db.execute(
                select(Participant).where(Participant.tenant_id == tenant_id, Participant.is_active.is_(True))
            ).scalars()
        )
        participants_by_id = {participant.id: participant for participant in participants}
        list_entries_by_list_id: dict[int, list[ListEntry]] = {}

        table_roles: dict[int, str] = {}
        table_list_definitions: dict[int, int | None] = {}
        table_matrix_keys: dict[int, str | None] = {}
        table_role_explicit: dict[int, bool] = {}
        for table in parsed.tables:

            def _list_entries_for(list_definition_id: int) -> list[ListEntry]:
                if list_definition_id not in list_entries_by_list_id:
                    list_entries_by_list_id[list_definition_id] = list(
                        db.execute(select(ListEntry).where(ListEntry.list_definition_id == list_definition_id)).scalars()
                    )
                return list_entries_by_list_id[list_definition_id]

            role, list_definition_id, matrix_key, explicit = _resolve_table_role(
                table, table_role_overrides or {}, table_roles_by_signature, list_definitions_for_matching, matrices_for_matching,
                list_definitions_by_id=list_definitions_by_id,
                list_entries_for=_list_entries_for,
                participants_by_id=participants_by_id,
            )
            table_roles[table.index] = role
            table_list_definitions[table.index] = list_definition_id
            table_matrix_keys[table.index] = matrix_key
            table_role_explicit[table.index] = explicit
        # Last-resort default when nothing in the document was recognized as the
        # attendance table at all: assume the first table is it. Must never override
        # table 0's role if that role came from an explicit source (this call's manual
        # override, or a signature learned from a previous import) - otherwise a
        # learned "table 0 is actually the Ämtli list, not attendance" mapping would
        # get silently clobbered back to "attendance" on every later import where the
        # real attendance table's heuristic match happens to miss. Also never overrides
        # a confident (non-explicit) "matrix" heading match - unlike "list", a matrix
        # match already names one specific, concrete template target (matrix_key), so
        # clobbering it is never desirable, whereas an unresolved "list" is allowed to
        # be reclaimed here (pre-existing behavior, left untouched).
        if (
            parsed.tables
            and not any(role == "attendance" for role in table_roles.values())
            and not table_role_explicit.get(parsed.tables[0].index, False)
            and table_roles.get(parsed.tables[0].index) != "matrix"
        ):
            table_roles[parsed.tables[0].index] = "attendance"

        form_type_id = db.scalar(select(ElementType.id).where(ElementType.code == "form"))
        template_linked_list_ids = _template_linked_list_ids(db, template_id=template_id, form_type_id=form_type_id)

        tables_preview = [
            TablePreview(
                index=table.index,
                header_cells=table.header_cells,
                sample_rows=table.rows[:3],
                role=table_roles.get(table.index, "ignore"),
                list_definition_id=table_list_definitions.get(table.index),
                matrix_key=table_matrix_keys.get(table.index),
                has_snapshot_target=(
                    table_list_definitions.get(table.index) in template_linked_list_ids
                    if table_roles.get(table.index) == "list"
                    else True
                ),
                role_is_explicit=table_role_explicit.get(table.index, False),
            )
            for table in parsed.tables
        ]
        tables_preview_by_index = {preview.index: preview for preview in tables_preview}

        text_type_id = db.scalar(select(ElementType.id).where(ElementType.code == "text"))
        static_text_type_id = db.scalar(select(ElementType.id).where(ElementType.code == "static_text"))
        text_capable_ids = {value for value in (text_type_id, static_text_type_id) if value is not None}
        # "form" blocks (fixed labeled rows, e.g. Organisation/Wer geht/Treffpunkt on a
        # Scharanlässe-style element) are offered as text targets too, alongside plain
        # text/static_text blocks - see is_form_block/form_rows below.
        content_capable_ids = text_capable_ids | ({form_type_id} if form_type_id is not None else set())

        template_rows = db.execute(
            select(TemplateElement, ElementDefinition)
            .join(ElementDefinition, ElementDefinition.id == TemplateElement.element_definition_id)
            .where(TemplateElement.template_id == template_id)
            .order_by(TemplateElement.sort_index.asc())
        ).all()

        text_targets: list[WordImportTextTarget] = []
        element_targets: dict[int, list[tuple[int, str]]] = {}
        element_titles: dict[int, str] = {}
        # Blocks whose repeat_source is "event" get one live instance per matching
        # Event (ProtocolService.add_event_block_to_element), not a fixed slot in the
        # template - a (template_element_id, block_sort_index) target alone can't tell
        # which instance/Event a section belongs to, so these keys are flagged here and
        # resolved to a specific Event further below, after the ordinary target
        # resolution (profile / title-similarity / manual) has picked the block exactly
        # like for any other element.
        event_repeat_block_keys: set[tuple[int, int]] = set()
        # (template_element_id, block_sort_index) -> that form block's raw configured rows
        # (id/label/row_type, straight from ElementDefinition), used below both to build
        # WordImportTextTarget.form_rows and, once a section is matched to such a target,
        # to extract per-row values out of its raw text via _parse_form_fields.
        form_rows_by_key: dict[tuple[int, int], list[dict]] = {}
        for template_element, definition in template_rows:
            blocks = sorted(
                (definition.configuration_json or {}).get("blocks", []),
                key=lambda entry: entry.get("sort_index", 0),
            )
            content_blocks = [block for block in blocks if block.get("element_type_id") in content_capable_ids]
            if not content_blocks:
                continue
            element_titles[template_element.id] = definition.title
            legacy_repeat_config = template_element.configuration_json or {}
            targets_for_element: list[tuple[int, str]] = []
            for block in content_blocks:
                sort_index = block["sort_index"]
                block_config = dict(block.get("configuration_json") or {})
                effective_repeat = (
                    {**legacy_repeat_config, **block_config}
                    if not block_config.get("repeat_source") and legacy_repeat_config.get("repeat_source")
                    else block_config
                )
                is_event_repeat_block = str(effective_repeat.get("repeat_source") or "") == "event"
                if is_event_repeat_block:
                    event_repeat_block_keys.add((template_element.id, sort_index))
                is_form_block = block.get("element_type_id") == form_type_id
                form_rows: list[dict] = []
                if is_form_block:
                    form_rows = sorted(block_config.get("rows") or [], key=lambda entry: entry.get("sort_index", 0))
                    form_rows_by_key[(template_element.id, sort_index)] = form_rows
                label = definition.title if len(content_blocks) == 1 else f'{definition.title} – {block.get("title") or f"Block {sort_index}"}'
                text_targets.append(
                    WordImportTextTarget(
                        template_element_id=template_element.id,
                        block_sort_index=sort_index,
                        label=label,
                        is_event_repeat=is_event_repeat_block,
                        is_form_block=is_form_block,
                        form_rows=[
                            WordImportFormRow(
                                row_id=str(row.get("id")),
                                label=str(row.get("label") or row.get("title") or "Feld"),
                                row_type=str(row.get("row_type") or row.get("value_type") or "text"),
                            )
                            for row in form_rows
                        ],
                    )
                )
                targets_for_element.append((sort_index, label))
            if targets_for_element:
                element_targets[template_element.id] = targets_for_element

        all_events = list(db.execute(select(Event).where(Event.tenant_id == tenant_id)).scalars())
        # Event-repeat blocks (e.g. "Rückblick", "Scharanlässe") must be linked to a real
        # Termin from the same Zyklus-Periode as the protocol being imported (the template's
        # configured CycleConfig, cycle year computed from protocol_date) - not just any
        # tenant event. Falls back to all tenant events if the template has no cycle
        # configured or no protocol_date could be determined yet.
        template = db.get(Template, template_id)
        cycle_cfg = db.get(CycleConfig, template.cycle_config_id) if template and template.cycle_config_id else None
        if cycle_cfg is not None and protocol_date is not None:
            target_cycle_year = get_cycle_year(protocol_date, cycle_cfg.reset_month, cycle_cfg.reset_day)
            period_events = [
                event
                for event in all_events
                if get_cycle_year(event.event_date, cycle_cfg.reset_month, cycle_cfg.reset_day) == target_cycle_year
            ]
        else:
            period_events = all_events

        text_mappings: list[WordImportTextMapping] = []
        for section in parsed.sections:
            normalized_heading = _normalize(section.heading)
            saved_target = heading_to_target.get(normalized_heading)
            if saved_target:
                template_element_id = saved_target.get("template_element_id")
                block_sort_index = saved_target.get("block_sort_index")
                confidence = 1.0
            else:
                best_element_id: int | None = None
                best_score = 0.0
                for candidate_element_id, title in element_titles.items():
                    if not element_targets.get(candidate_element_id):
                        continue
                    score = _similarity(section.heading, title)
                    if score > best_score:
                        best_score = score
                        best_element_id = candidate_element_id
                if best_element_id is not None and best_score >= 0.35:
                    template_element_id = best_element_id
                    block_sort_index = min(sort_index for sort_index, _ in element_targets[best_element_id])
                    confidence = round(best_score, 2)
                else:
                    template_element_id = None
                    block_sort_index = None
                    confidence = round(best_score, 2)
                    warnings.append(f'Kein passendes Element für Abschnitt "{section.heading}" gefunden – bitte manuell zuweisen.')

            is_event_repeat = (
                template_element_id is not None
                and block_sort_index is not None
                and (template_element_id, block_sort_index) in event_repeat_block_keys
            )
            # Computed unconditionally (not gated on is_event_repeat) so that manually
            # switching a section's target to a DIFFERENT event-repeat block later in the
            # wizard - after this initial auto-match - still has a full, non-empty
            # candidate list to show instead of the empty [] a mapping starts with. Since
            # this list is period-scoped (not per-block tag/window-filtered), it's valid
            # for every event-repeat target in this protocol, not just the auto-matched one.
            search_text = _strip_title_prefix(
                section.heading, element_titles.get(template_element_id, "") if template_element_id is not None else ""
            )
            scored_events = sorted(
                ((_similarity(search_text, event.title), event) for event in period_events),
                key=lambda entry: entry[0],
                reverse=True,
            )
            event_candidates = [
                WordImportEventCandidate(
                    event_id=event.id, title=event.title, event_date=event.event_date, score=round(score, 3),
                    reason=_text_match_reason(search_text, event.title, score),
                )
                for score, event in scored_events
            ]
            matched_event_id: int | None = None
            if is_event_repeat:
                best_event_score, best_event = (scored_events[0] if scored_events else (0.0, None))
                matched_event_id = best_event.id if best_event is not None and best_event_score >= _EVENT_REPEAT_MATCH_THRESHOLD else None
                if matched_event_id is None:
                    warnings.append(f'Rückblick "{section.heading}" konnte keinem Anlass zugeordnet werden – bitte manuell wählen.')

            is_form_block = (
                template_element_id is not None
                and block_sort_index is not None
                and (template_element_id, block_sort_index) in form_rows_by_key
            )
            # Parsed against EVERY form-block target (not just the one currently matched
            # above) so that manually switching a section's target later in the wizard -
            # e.g. because the initial title-similarity match missed and
            # template_element_id started out None - still has real parsed values to show
            # instead of silently falling back to blank fields. Same reasoning as
            # event_candidates being computed unconditionally further up.
            form_fields_by_target: dict[str, list[WordImportFormFieldValue]] = {}
            for form_key, rows in form_rows_by_key.items():
                parsed_field_values = _parse_form_fields(section.text, rows)
                fields_for_target: list[WordImportFormFieldValue] = []
                for row in rows:
                    row_id = str(row.get("id"))
                    row_type = str(row.get("row_type") or row.get("value_type") or "text")
                    raw_value = parsed_field_values.get(row_id, "")
                    names: list[WordImportNameResolution] = []
                    if row_type in ("participant", "participants") and raw_value:
                        names = _match_names(raw_value, participants, participant_name_overrides, rejected_candidates, participant_match_threshold)
                    fields_for_target.append(
                        WordImportFormFieldValue(
                            row_id=row_id,
                            label=str(row.get("label") or "Feld"),
                            row_type=row_type,
                            raw_value=raw_value,
                            names=names,
                        )
                    )
                form_fields_by_target[f"{form_key[0]}:{form_key[1]}"] = fields_for_target
                if is_form_block and form_key == (template_element_id, block_sort_index) and not parsed_field_values:
                    warnings.append(f'Für "{section.heading}" konnten keine Felder erkannt werden – bitte manuell ausfüllen.')
            form_fields = (
                form_fields_by_target.get(f"{template_element_id}:{block_sort_index}", []) if is_form_block else []
            )

            text_mappings.append(
                WordImportTextMapping(
                    extracted_heading=section.heading,
                    extracted_text=section.text,
                    template_element_id=template_element_id,
                    block_sort_index=block_sort_index,
                    confidence=confidence,
                    is_event_repeat=is_event_repeat,
                    matched_event_id=matched_event_id,
                    event_candidates=event_candidates,
                    is_form_block=is_form_block,
                    form_fields=form_fields,
                    form_fields_by_target=form_fields_by_target,
                )
            )

        # Attendance name-matching must only consider participants who actually belong to
        # this template's attendance roster, not every active participant in the tenant -
        # otherwise a raw name like "Dominik Rohrer" with no roster match can fuzzy-match
        # onto an unrelated "Armin Rohrer" from another template purely by shared surname
        # (SequenceMatcher scores full-name overlap, so a shared last name alone can clear
        # _PARTICIPANT_MATCH_THRESHOLD even with a completely different first name).
        roster_filters = [
            TemplateParticipant.template_id == template_id,
            TemplateParticipant.exclude_from_attendance.is_(False),
        ]
        if protocol_date is not None:
            roster_filters.append(participant_eligible_on(protocol_date))
        template_roster = list(
            db.execute(
                select(Participant)
                .join(TemplateParticipant, TemplateParticipant.participant_id == Participant.id)
                .where(*roster_filters)
                .order_by(Participant.display_name.asc(), Participant.id.asc())
            ).scalars()
        )

        raw_attendance_rows: list[tuple[str, str]] = []
        for table in parsed.tables:
            if table_roles.get(table.index) != "attendance":
                continue
            for cells in table.rows:
                if not cells or not cells[0]:
                    continue
                status = "present"
                for marker in cells[1:]:
                    classified = _classify_status(marker)
                    if classified:
                        status = classified
                        break
                raw_attendance_rows.append((cells[0], status))

        # Two document rows must never end up auto-suggested to the same participant -
        # once a participant has been claimed, later rows can no longer match them, so
        # they either fall to a weaker candidate or stay unmatched for a human to resolve.
        used_participant_ids: set[int] = set()
        row_assignment: dict[int, int] = {}

        # Saved overrides are claimed first, in document order, regardless of fuzzy score.
        for row_index, (raw_name, _status) in enumerate(raw_attendance_rows):
            override_id = participant_name_overrides.get(_normalize(raw_name))
            if override_id is not None and override_id not in used_participant_ids:
                row_assignment[row_index] = override_id
                used_participant_ids.add(override_id)

        # Remaining rows must be resolved by GLOBAL best score first, not in document row
        # order - otherwise an early row with only a mediocre same-surname match (e.g.
        # "Dominik Rohrer" fuzzy-matching "Mario Rohrer" at 0.77, since SequenceMatcher
        # scores full-name overlap and a shared surname alone can clear the threshold)
        # would claim that participant before a later row with the actual exact match
        # ("Mario Rohrer" == "Mario Rohrer", 1.0) ever gets a chance - leaving the correct
        # row stranded on "Keinen verknüpfen" even though a perfect match existed.
        def _attendance_score(row_index: int, participant: Participant) -> float:
            raw_name = raw_attendance_rows[row_index][0]
            score = _name_score(raw_name, participant.display_name)
            if participant.id in _rejected_ids_for(f"name:{_normalize(raw_name)}"):
                score -= _REJECTED_CANDIDATE_PENALTY
            return score

        remaining_row_indices = [
            row_index for row_index in range(len(raw_attendance_rows)) if row_index not in row_assignment
        ]
        remaining_roster = [participant for participant in template_roster if participant.id not in used_participant_ids]
        # Globally optimal assignment - see _match_names's identical use of
        # solve_optimal_assignment for why this replaces the previous greedy pass.
        for assignment in solve_optimal_assignment(
            remaining_row_indices, remaining_roster, _attendance_score, min_score=participant_match_threshold
        ):
            row_assignment[assignment.row] = assignment.col.id
            used_participant_ids.add(assignment.col.id)

        attendance_mappings: list[WordImportAttendanceMapping] = []
        for row_index, (raw_name, status) in enumerate(raw_attendance_rows):
            suggested = row_assignment.get(row_index)
            if suggested is None:
                warnings.append(f'Kein passender Teilnehmer für "{raw_name}" gefunden.')
            candidates = [
                WordImportAttendanceCandidate(
                    participant_id=participant.id, score=round(score, 3),
                    reason=_text_match_reason(raw_name, participant.display_name, score),
                )
                for score, participant in sorted(
                    ((_name_score(raw_name, participant.display_name), participant) for participant in template_roster),
                    key=lambda entry: entry[0],
                    reverse=True,
                )[:3]
                if score >= 0.4
            ]
            attendance_mappings.append(
                WordImportAttendanceMapping(
                    raw_name=raw_name, status=status, suggested_participant_id=suggested, candidates=candidates
                )
            )

        # Participants who belong to this template's attendance roster (same query
        # ProtocolService.create_from_template uses to build attendance_entries) but were
        # never mentioned by name anywhere in the document must still show up in the
        # review step - otherwise they'd silently default to "absent" with no visible row,
        # looking like missing data instead of an explicit, editable default.
        already_matched_participant_ids = {
            mapping.suggested_participant_id for mapping in attendance_mappings if mapping.suggested_participant_id is not None
        }
        for participant in template_roster:
            if participant.id in already_matched_participant_ids:
                continue
            attendance_mappings.append(
                WordImportAttendanceMapping(
                    raw_name="", status="absent", suggested_participant_id=participant.id,
                    candidates=[
                        WordImportAttendanceCandidate(
                            participant_id=participant.id, score=1.0,
                            reason="Nicht im Dokument erwähnt, aus Vorlagen-Roster übernommen",
                        )
                    ],
                )
            )

        event_mappings: list[WordImportEventMapping] = []
        # Collected first (title, raw_date, non-exclusive candidate list) per row across
        # every "events"-role table, THEN resolved against all_events in one single
        # optimal assignment below - not per-row independently anymore. Without this,
        # two different document rows could both auto-match the SAME existing Event
        # (verified: nothing previously stopped that), which is almost always wrong for
        # an events/Termine table, where each row names a distinct real occasion.
        extracted_event_rows: list[tuple[str, date | None, list[WordImportEventCandidate]]] = []
        for table in parsed.tables:
            if table_roles.get(table.index) != "events":
                continue
            for cells in table.rows:
                title, raw_date = _extract_event_row(cells)
                if not title:
                    continue
                event_rejected_ids = (
                    _rejected_ids_for(f"event:{raw_date.isoformat()}|{_normalize(title)}") if raw_date is not None else set()
                )
                scored_events = sorted(
                    (
                        (
                            _score_event_candidate(title, raw_date, event)
                            - (_REJECTED_CANDIDATE_PENALTY if event.id in event_rejected_ids else 0.0),
                            event,
                        )
                        for event in all_events
                    ),
                    key=lambda entry: entry[0],
                    reverse=True,
                )
                candidates = [
                    WordImportEventCandidate(
                        event_id=event.id, title=event.title, event_date=event.event_date, score=round(score, 3),
                        reason=_event_match_reason(title, raw_date, event),
                    )
                    for score, event in scored_events[:_CANDIDATE_LIMIT]
                ]
                extracted_event_rows.append((title, raw_date, candidates))

        def _event_row_score(row_index: int, event: Event) -> float:
            title, raw_date, _candidates = extracted_event_rows[row_index]
            rejected_ids = (
                _rejected_ids_for(f"event:{raw_date.isoformat()}|{_normalize(title)}") if raw_date is not None else set()
            )
            score = _score_event_candidate(title, raw_date, event)
            return score - (_REJECTED_CANDIDATE_PENALTY if event.id in rejected_ids else 0.0)

        # min_score is the loose "changed" threshold - it decides "is this row linked to
        # an existing event AT ALL", not the finer matched-vs-changed distinction (that's
        # still re-derived per row below, exactly as before, just against the one
        # globally-assigned event instead of each row's own locally-best one).
        auto_picked_event_by_row: dict[int, Event] = {
            assignment.row: assignment.col
            for assignment in solve_optimal_assignment(
                list(range(len(extracted_event_rows))), all_events, _event_row_score, min_score=event_change_threshold
            )
        }
        for row_index, (title, raw_date, candidates) in enumerate(extracted_event_rows):
            best_event = auto_picked_event_by_row.get(row_index)
            if best_event is not None and raw_date == best_event.event_date and _similarity(title, best_event.title) >= _EVENT_MATCH_THRESHOLD:
                status: str = "matched"
            elif best_event is not None:
                status = "changed"
            else:
                status = "new"
            event_mappings.append(
                WordImportEventMapping(
                    row_index=row_index,
                    raw_title=title,
                    raw_date=raw_date,
                    status=status,
                    matched_event_id=best_event.id if best_event else None,
                    matched_event_title=best_event.title if best_event else None,
                    matched_event_date=best_event.event_date if best_event else None,
                    candidates=candidates,
                )
            )
        # The Matrix "events" row section further below continues this same counter
        # (see its own comment there) - explicit assignment rather than relying on the
        # loop variable's last value, which would incorrectly equal len(...)-1, not
        # len(...), once the for-loop above completes.
        row_index = len(extracted_event_rows)
        if any(mapping.status != "matched" for mapping in event_mappings):
            warnings.append("Neue oder abweichende Termine gefunden – bitte prüfen und übernehmen oder ablehnen.")

        previous_protocol_id = (
            _nearest_previous_protocol_id(db, tenant_id=tenant_id, template_id=template_id, protocol_date=protocol_date)
            if protocol_date is not None
            else None
        )
        previous_entries_by_list: dict[int, dict[int, dict]] = {}

        list_mappings: list[WordImportListRowMapping] = []
        for table in parsed.tables:
            if table_roles.get(table.index) != "list":
                continue
            list_definition_id = table_list_definitions.get(table.index)
            if list_definition_id is None or list_definition_id not in list_definitions_by_id:
                warnings.append(f"Tabelle #{table.index + 1}: keine passende Liste gefunden – bitte manuell auswählen.")
                continue
            has_snapshot_target = list_definition_id in template_linked_list_ids
            if not has_snapshot_target:
                warnings.append(
                    f'Vorlage hat keinen Block für Liste "{list_definitions_by_id[list_definition_id].name}" – '
                    "Tabelle wird nicht importiert."
                )
            definition = list_definitions_by_id[list_definition_id]
            existing_entries = list(
                db.execute(select(ListEntry).where(ListEntry.list_definition_id == list_definition_id)).scalars()
            )
            existing_by_key = {
                _value_key(definition.column_one_value_type, entry.column_one_value_json or {}): entry
                for entry in existing_entries
            }
            if list_definition_id not in previous_entries_by_list:
                previous_entries_by_list[list_definition_id] = (
                    _list_snapshot_entries_by_id(db, protocol_id=previous_protocol_id, list_definition_id=list_definition_id)
                    if previous_protocol_id is not None
                    else {}
                )
            previous_entries = previous_entries_by_list[list_definition_id]
            table_signature = _normalize(" | ".join(table.header_cells))
            forced_strategy = (table_role_overrides or {}).get(table.index, {}).get(
                "list_grouping_strategy"
            ) or table_roles_by_signature.get(table_signature, {}).get("list_grouping_strategy")
            grouping_strategy, row_candidates, needs_manual_grouping, available_grouping_strategies = _select_list_row_variant(
                table.rows, definition, existing_entries, participants_by_id, forced_strategy
            )
            table_preview = tables_preview_by_index.get(table.index)
            if table_preview is not None:
                table_preview.grouping_strategy = grouping_strategy
                table_preview.needs_manual_grouping = needs_manual_grouping
                table_preview.available_grouping_strategies = available_grouping_strategies
            for list_row_index, row_candidate in enumerate(row_candidates):
                column_one_raw = row_candidate.column_one_raw
                column_two_raw = row_candidate.column_two_raw
                col1_value, col1_names = _build_column_value(
                    definition.column_one_value_type, column_one_raw, participants, participant_name_overrides, rejected_candidates,
                    participant_match_threshold,
                )
                col2_value, col2_names = _build_column_value(
                    definition.column_two_value_type, column_two_raw, participants, participant_name_overrides, rejected_candidates,
                    participant_match_threshold,
                )
                key = _value_key(definition.column_one_value_type, col1_value)
                existing = existing_by_key.get(key)
                scored_entries = sorted(
                    (
                        (
                            _similarity(
                                column_one_raw,
                                _display_value(definition.column_one_value_type, entry.column_one_value_json or {}, participants_by_id),
                            ),
                            entry,
                        )
                        for entry in existing_entries
                    ),
                    key=lambda item: item[0],
                    reverse=True,
                )
                candidates = [
                    WordImportListEntryCandidate(
                        entry_id=entry.id,
                        column_one_display=_display_value(definition.column_one_value_type, entry.column_one_value_json or {}, participants_by_id),
                        column_two_display=_display_value(definition.column_two_value_type, entry.column_two_value_json or {}, participants_by_id),
                        score=round(score, 3),
                        reason=_text_match_reason(
                            column_one_raw,
                            _display_value(definition.column_one_value_type, entry.column_one_value_json or {}, participants_by_id),
                            score,
                        ),
                    )
                    for score, entry in scored_entries
                    if score >= _LIST_ENTRY_CANDIDATE_MIN_SCORE
                ][:_CANDIDATE_LIMIT]
                if existing is not None and not any(candidate.entry_id == existing.id for candidate in candidates):
                    candidates.insert(
                        0,
                        WordImportListEntryCandidate(
                            entry_id=existing.id,
                            column_one_display=_display_value(definition.column_one_value_type, existing.column_one_value_json or {}, participants_by_id),
                            column_two_display=_display_value(definition.column_two_value_type, existing.column_two_value_json or {}, participants_by_id),
                            score=1.0,
                            reason="Bereits mit dieser Dokumentzeile verknüpfter Eintrag",
                        ),
                    )
                if existing is None:
                    status = "new"
                    matched_entry_id = None
                else:
                    # Diff against the nearest earlier protocol's frozen snapshot value for
                    # this same entry when available (a list drifts over time - comparing an
                    # old document to *today's* live value would show spurious "changed"
                    # rows for entries nothing actually changed about back then), otherwise
                    # fall back to the live value.
                    baseline_col2 = (previous_entries.get(existing.id) or {}).get("column_two_value")
                    if baseline_col2 is None:
                        baseline_col2 = existing.column_two_value_json or {}
                    baseline_col2_key = _value_key(definition.column_two_value_type, baseline_col2)
                    new_col2_key = _value_key(definition.column_two_value_type, col2_value)
                    status = "matched" if baseline_col2_key == new_col2_key else "changed"
                    matched_entry_id = existing.id
                unmatched_names = [
                    resolution.raw_name for resolution in (col1_names + col2_names) if resolution.participant_id is None
                ]
                if unmatched_names:
                    warnings.append(f'Nicht gefundene Teilnehmer in Liste "{definition.name}": {", ".join(unmatched_names)}')
                list_mappings.append(
                    WordImportListRowMapping(
                        table_index=table.index,
                        row_index=list_row_index,
                        column_one_raw=column_one_raw,
                        column_two_raw=column_two_raw,
                        column_one_type=definition.column_one_value_type,
                        column_two_type=definition.column_two_value_type,
                        status=status,
                        matched_entry_id=matched_entry_id,
                        column_one_names=col1_names,
                        column_two_names=col2_names,
                        candidates=candidates,
                        has_snapshot_target=has_snapshot_target,
                        group_filled=row_candidate.group_filled,
                    )
                )

        matrix_mappings: list[WordImportMatrixCellMapping] = []
        # Reuses the same lazy per-list-id cache the list-vs-matrix role tie-break
        # above already declared and may have partially populated - not redeclared here.
        for table in parsed.tables:
            if table_roles.get(table.index) != "matrix":
                continue
            matrix_key = table_matrix_keys.get(table.index)
            matrix = matrices_by_key.get(matrix_key) if matrix_key else None
            if matrix is None:
                warnings.append(f"Tabelle #{table.index + 1}: keine passende Matrix gefunden – bitte manuell auswählen.")
                continue
            matrix_rows = matrix["rows"]
            matrix_columns = matrix["columns"]
            mode = matrix["mode"]
            auto_source = matrix["auto_source"] or {}
            auto_type = str(auto_source.get("type") or "")
            # First row (minus the corner cell) = column headers in the document; each
            # data row's cells[0] is the row label, cells[1:] align positionally with
            # these headers - the "classic cross-table" shape confirmed with Timo.
            doc_column_labels = table.header_cells[1:]

            def _entry_title(entry: ListEntry) -> str:
                col1 = entry.column_one_value_json if isinstance(entry.column_one_value_json, dict) else {}
                col2 = entry.column_two_value_json if isinstance(entry.column_two_value_json, dict) else {}
                return str(col1.get("text_value") or col2.get("text_value") or f"Eintrag {entry.id}")

            # Column resolution is computed once per table (shared across all rows),
            # same idea as row resolution below - column_resolution[col_idx] = (chosen
            # column_key or None, ranked candidates for a manual pick).
            column_resolution: dict[int, tuple[str | None, list[WordImportMatrixColumnCandidate]]] = {}
            if mode == "auto" and auto_type == "participants":
                # Two sub-passes so solve_optimal_assignment's exclusivity only applies
                # where columns actually compete for the same pool: a column resolved by
                # a direct name match (_match_names, itself already Hungarian-backed
                # internally) claims its participant immediately, then the remaining
                # ambiguous columns are assigned optimally over the still-unclaimed
                # participants - never re-opening an already-confident direct match.
                unresolved_col_indices: list[int] = []
                claimed_participant_ids: set[int] = set()
                for col_idx, label in enumerate(doc_column_labels):
                    if not label:
                        continue
                    names = _match_names(label, participants, participant_name_overrides, rejected_candidates, participant_match_threshold)
                    participant_id = names[0].participant_id if names else None
                    if participant_id is not None:
                        participant = next((p for p in participants if p.id == participant_id), None)
                        column_key = f"gen-p-{participant_id}"
                        column_resolution[col_idx] = (
                            column_key,
                            [
                                WordImportMatrixColumnCandidate(
                                    column_key=column_key, label=participant.display_name if participant else label, score=1.0,
                                    reason="Direkter Namenstreffer",
                                )
                            ],
                        )
                        claimed_participant_ids.add(participant_id)
                    else:
                        unresolved_col_indices.append(col_idx)

                def _participant_column_score(col_idx: int, participant: Participant) -> float:
                    return _similarity(doc_column_labels[col_idx], participant.display_name)

                remaining_participants = [p for p in participants if p.id not in claimed_participant_ids]
                auto_picked_participant_by_col = {
                    assignment.row: assignment.col
                    for assignment in solve_optimal_assignment(
                        unresolved_col_indices, remaining_participants, _participant_column_score,
                        min_score=_MATRIX_COLUMN_MATCH_THRESHOLD,
                    )
                }
                for col_idx in unresolved_col_indices:
                    label = doc_column_labels[col_idx]
                    scored = sorted(
                        ((_similarity(label, p.display_name), p) for p in participants),
                        key=lambda entry: entry[0],
                        reverse=True,
                    )[:_CANDIDATE_LIMIT]
                    picked = auto_picked_participant_by_col.get(col_idx)
                    column_resolution[col_idx] = (
                        f"gen-p-{picked.id}" if picked is not None else None,
                        [
                            WordImportMatrixColumnCandidate(
                                column_key=f"gen-p-{p.id}", label=p.display_name, score=round(score, 3),
                                reason=_text_match_reason(label, p.display_name, score),
                            )
                            for score, p in scored
                            if score >= _LIST_ENTRY_CANDIDATE_MIN_SCORE
                        ],
                    )
            elif mode == "auto" and auto_type == "events":
                tag_filter = str(auto_source.get("event_tag_filter") or "").strip().lower()
                candidate_events = [event for event in all_events if not tag_filter or (event.tag or "").lower() == tag_filter]
                labeled_col_indices = [col_idx for col_idx, label in enumerate(doc_column_labels) if label]

                def _event_column_score(col_idx: int, event: Event) -> float:
                    return _similarity(doc_column_labels[col_idx], event.title)

                # Cross-column exclusivity is a genuine behavior change from before (each
                # column used to independently pick its own best event with no regard for
                # what another column already claimed) - two columns claiming the same
                # Event is almost never correct for a cross-table Matrix, so this is
                # desirable, mirroring the same change made to Termine-table row matching.
                auto_picked_event_by_col = {
                    assignment.row: assignment.col
                    for assignment in solve_optimal_assignment(
                        labeled_col_indices, candidate_events, _event_column_score, min_score=_MATRIX_COLUMN_MATCH_THRESHOLD
                    )
                }
                for col_idx in labeled_col_indices:
                    label = doc_column_labels[col_idx]
                    scored = sorted(
                        ((_similarity(label, event.title), event) for event in candidate_events),
                        key=lambda entry: entry[0],
                        reverse=True,
                    )
                    picked = auto_picked_event_by_col.get(col_idx)
                    column_resolution[col_idx] = (
                        f"gen-e-{picked.id}" if picked is not None else None,
                        [
                            WordImportMatrixColumnCandidate(
                                column_key=f"gen-e-{event.id}", label=event.title, score=round(score, 3),
                                reason=_text_match_reason(label, event.title, score),
                            )
                            for score, event in scored[:_CANDIDATE_LIMIT]
                            if score >= _LIST_ENTRY_CANDIDATE_MIN_SCORE
                        ],
                    )
            elif mode == "auto" and auto_type == "list":
                list_id = auto_source.get("list_id")
                entries: list[ListEntry] = []
                if list_id:
                    list_id = int(list_id)
                    if list_id not in list_entries_by_list_id:
                        list_entries_by_list_id[list_id] = list(
                            db.execute(select(ListEntry).where(ListEntry.list_definition_id == list_id)).scalars()
                        )
                    entries = list_entries_by_list_id[list_id]
                labeled_col_indices = [col_idx for col_idx, label in enumerate(doc_column_labels) if label]

                def _list_column_score(col_idx: int, entry: ListEntry) -> float:
                    return _similarity(doc_column_labels[col_idx], _entry_title(entry))

                # Same exclusivity rationale as the "events" branch above.
                auto_picked_entry_by_col = {
                    assignment.row: assignment.col
                    for assignment in solve_optimal_assignment(
                        labeled_col_indices, entries, _list_column_score, min_score=_MATRIX_COLUMN_MATCH_THRESHOLD
                    )
                }
                for col_idx in labeled_col_indices:
                    label = doc_column_labels[col_idx]
                    scored = sorted(
                        ((_similarity(label, _entry_title(entry)), entry) for entry in entries),
                        key=lambda item: item[0],
                        reverse=True,
                    )
                    picked = auto_picked_entry_by_col.get(col_idx)
                    column_resolution[col_idx] = (
                        f"gen-l-{picked.id}" if picked is not None else None,
                        [
                            WordImportMatrixColumnCandidate(
                                column_key=f"gen-l-{entry.id}", label=_entry_title(entry), score=round(score, 3),
                                reason=_text_match_reason(label, _entry_title(entry), score),
                            )
                            for score, entry in scored[:_CANDIDATE_LIMIT]
                            if score >= _LIST_ENTRY_CANDIDATE_MIN_SCORE
                        ],
                    )
            else:
                for col_idx, label in enumerate(doc_column_labels):
                    if not label:
                        continue
                    scored = sorted(
                        ((_similarity(label, str(column.get("title") or "")), column) for column in matrix_columns),
                        key=lambda entry: entry[0],
                        reverse=True,
                    )
                    best_score, best_column = scored[0] if scored else (0.0, None)
                    column_key = (
                        str(best_column.get("id")) if best_column is not None and best_score >= _MATRIX_COLUMN_MATCH_THRESHOLD else None
                    )
                    column_resolution[col_idx] = (
                        column_key,
                        [
                            WordImportMatrixColumnCandidate(
                                column_key=str(column.get("id")), label=str(column.get("title") or ""), score=round(score, 3),
                                reason=_text_match_reason(label, str(column.get("title") or ""), score),
                            )
                            for score, column in scored[:_CANDIDATE_LIMIT]
                            if score >= _LIST_ENTRY_CANDIDATE_MIN_SCORE
                        ],
                    )
                # Positional fallback (see B.6(b) / _positional_matrix_column_resolution
                # docstring) - only tried when label-similarity matching resolved fewer
                # than half of this table's labeled columns, and only kept if it actually
                # resolves MORE than the label-based attempt did. Guards against ever
                # silently overriding an already-adequate label match with a guess.
                labeled_col_count = sum(1 for label in doc_column_labels if label)
                if labeled_col_count > 0 and _count_confident_matrix_columns(column_resolution) < labeled_col_count / 2:
                    positional_resolution = _positional_matrix_column_resolution(doc_column_labels, matrix_columns)
                    if _count_confident_matrix_columns(positional_resolution) > _count_confident_matrix_columns(column_resolution):
                        column_resolution = positional_resolution
                        warnings.append(
                            f'Matrix "{matrix["title"]}": Spalten konnten über die Bezeichnung kaum zugeordnet werden – '
                            "stattdessen anhand der Position übernommen, bitte prüfen."
                        )

            matrix_has_unresolved_column = False
            for row_cells in table.rows:
                row_label_raw = row_cells[0] if row_cells else ""
                if not row_label_raw:
                    continue
                scored_rows = sorted(
                    ((_similarity(row_label_raw, str(row.get("label") or row.get("title") or "")), row) for row in matrix_rows),
                    key=lambda entry: entry[0],
                    reverse=True,
                )
                best_row_score, best_row = scored_rows[0] if scored_rows else (0.0, None)
                if best_row is None or best_row_score < _MATRIX_ROW_MATCH_THRESHOLD:
                    warnings.append(f'Matrix "{matrix["title"]}": Zeile "{row_label_raw}" konnte keiner Zeile zugeordnet werden – wird übersprungen.')
                    continue
                row_type = str(best_row.get("row_type") or "text")
                # row_type "7" is an embedded Terminliste block (element_type_id 7) - the
                # far more common way this is actually authored in practice (confirmed
                # against Timo's real template) than the plain named "events" row_type.
                # Both render through the exact same export_service._matrix_event_row_value
                # (an embedded cell only ever overrides that with its own stored
                # `embedded_block` content if one was manually edited in a specific live
                # protocol - never the case for a freshly imported one) and share the same
                # row_config shape (event_tag_filter/event_use_column_tag_filter/etc.), so
                # both are handled identically here.
                if row_type in ("events", "7"):
                    # A Matrix "events" row never stores its dates per cell - they're
                    # resolved live at render time by matching an Event's own `tag`
                    # against the column (see export_service._matrix_events). So there is
                    # nothing to write into row_values here; instead each date found in a
                    # cell is folded into the SAME event_mappings/Termine review used for
                    # ordinary "events"-role tables (continuing the same row_index
                    # counter), carrying the tag this Event needs so it actually shows up
                    # in this Matrix column once committed.
                    row_config = best_row.get("row_config") or {}
                    use_column_tag = bool(
                        best_row.get("event_use_column_tag_filter")
                        or row_config.get("event_use_column_tag_filter")
                        or best_row.get("use_column_title_as_tag", row_config.get("use_column_title_as_tag", True))
                    )
                    for col_idx, column_label_raw in enumerate(doc_column_labels):
                        if not column_label_raw or col_idx + 1 >= len(row_cells):
                            continue
                        raw_cell_text = row_cells[col_idx + 1]
                        if not raw_cell_text:
                            continue
                        column_key, _column_candidates = column_resolution.get(col_idx, (None, []))
                        if column_key is None:
                            matrix_has_unresolved_column = True
                            warnings.append(
                                f'Matrix "{matrix["title"]}": Spalte "{column_label_raw}" konnte nicht zugeordnet werden – '
                                "Termine aus dieser Spalte werden übersprungen."
                            )
                            continue
                        matched_column = next((column for column in matrix_columns if str(column.get("id")) == column_key), None)
                        column_title = str((matched_column or {}).get("title") or column_label_raw)
                        effective_tag = None
                        if use_column_tag and matched_column is not None:
                            effective_tag = str(matched_column.get("event_tag_filter") or matched_column.get("title") or "").strip() or None
                        for extracted_date, extracted_participant_count in _extract_dates_with_counts(raw_cell_text):
                            matrix_event_rejected_ids = _rejected_ids_for(
                                f"event:{extracted_date.isoformat()}|{_normalize(column_title)}"
                            )
                            scored_events = sorted(
                                (
                                    (
                                        _score_event_candidate(column_title, extracted_date, event)
                                        - (_REJECTED_CANDIDATE_PENALTY if event.id in matrix_event_rejected_ids else 0.0),
                                        event,
                                    )
                                    for event in all_events
                                ),
                                key=lambda entry: entry[0],
                                reverse=True,
                            )
                            candidates = [
                                WordImportEventCandidate(
                                    event_id=event.id, title=event.title, event_date=event.event_date, score=round(score, 3),
                                    reason=_event_match_reason(column_title, extracted_date, event),
                                )
                                for score, event in scored_events[:_CANDIDATE_LIMIT]
                            ]
                            # Matrix cells never carry a real title to compare - a plain
                            # exact date match is the only meaningful signal here (unlike
                            # ordinary Termine rows, whose title-similarity gate a bare
                            # date coincidence would too easily satisfy by chance).
                            exact_date_match = next((event for event in all_events if event.event_date == extracted_date), None)
                            matched_event = exact_date_match
                            status = "matched" if matched_event is not None else "new"
                            event_mappings.append(
                                WordImportEventMapping(
                                    row_index=row_index,
                                    raw_title=column_title,
                                    raw_date=extracted_date,
                                    status=status,
                                    matched_event_id=matched_event.id if matched_event else None,
                                    matched_event_title=matched_event.title if matched_event else None,
                                    matched_event_date=matched_event.event_date if matched_event else None,
                                    candidates=candidates,
                                    tag=effective_tag,
                                    participant_count=extracted_participant_count,
                                    matrix_key=matrix_key,
                                    matrix_title=matrix["title"],
                                    row_id=str(best_row.get("id")),
                                    row_label=str(best_row.get("label") or best_row.get("title") or ""),
                                    column_key=column_key,
                                    column_label=column_title,
                                )
                            )
                            row_index += 1
                    continue
                if row_type not in _MATRIX_SUPPORTED_ROW_TYPES:
                    warnings.append(f'Matrix "{matrix["title"]}": Zeile "{row_label_raw}" hat einen nicht unterstützten Zeilentyp und wird übersprungen.')
                    continue
                row_id = str(best_row.get("id"))
                for col_idx, column_label_raw in enumerate(doc_column_labels):
                    if not column_label_raw or col_idx + 1 >= len(row_cells):
                        continue
                    raw_value = row_cells[col_idx + 1]
                    if not raw_value:
                        continue
                    column_key, column_candidates = column_resolution.get(col_idx, (None, []))
                    if column_key is None:
                        matrix_has_unresolved_column = True
                    names: list[WordImportNameResolution] = []
                    if row_type in ("participant", "participants"):
                        names = _match_names(raw_value, participants, participant_name_overrides, rejected_candidates, participant_match_threshold)
                        unmatched_names = [name.raw_name for name in names if name.participant_id is None]
                        if unmatched_names:
                            warnings.append(f'Nicht gefundene Teilnehmer in Matrix "{matrix["title"]}": {", ".join(unmatched_names)}')
                    matrix_mappings.append(
                        WordImportMatrixCellMapping(
                            table_index=table.index,
                            matrix_key=matrix_key,
                            matrix_title=matrix["title"],
                            row_id=row_id,
                            row_label=str(best_row.get("label") or best_row.get("title") or ""),
                            row_label_raw=row_label_raw,
                            row_type=row_type,
                            column_label_raw=column_label_raw,
                            column_key=column_key,
                            column_candidates=column_candidates,
                            raw_value=raw_value,
                            names=names,
                        )
                    )
            if matrix_has_unresolved_column:
                warnings.append(f'Matrix "{matrix["title"]}": nicht alle Spalten konnten eindeutig zugeordnet werden – bitte prüfen.')

        # The same event can legitimately be extracted twice - e.g. a duplicate row in
        # the document's own "Termine" table, or (less commonly) the same date showing
        # up in two Matrix cells of the same column - which would otherwise show the
        # identical "Titel (Datum)" entry twice in the Termine review and let the user
        # accidentally create two duplicate Events for the same occasion. Keep only the
        # first occurrence per (title, date, matrix column) triple.
        seen_event_keys: set[tuple[str, date | None, str | None, str | None]] = set()
        deduped_event_mappings: list[WordImportEventMapping] = []
        for mapping in event_mappings:
            key = (_normalize(mapping.raw_title), mapping.raw_date, mapping.matrix_key, mapping.column_key)
            if key in seen_event_keys:
                continue
            seen_event_keys.add(key)
            deduped_event_mappings.append(mapping)
        event_mappings = deduped_event_mappings

        if protocol_date is None:
            warnings.append("Kein Datum im Dokument erkannt – bitte manuell angeben.")

        return WordImportAnalysis(
            protocol_date=protocol_date,
            tables=tables_preview,
            text_mappings=text_mappings,
            text_targets=text_targets,
            attendance_mappings=attendance_mappings,
            event_mappings=event_mappings,
            list_definitions=[WordImportListDefinitionOption(id=item.id, name=item.name) for item in list_definition_rows],
            list_mappings=list_mappings,
            matrix_options=[WordImportMatrixOption(matrix_key=matrix["matrix_key"], title=matrix["title"]) for matrix in matrices_for_matching],
            matrix_mappings=matrix_mappings,
            profile_applied=profile_applied,
            warnings=warnings,
        )

    def commit(self, db: Session, *, tenant_id: int, user_id: int | None, payload: WordImportCommit) -> int:
        protocol_service = ProtocolService()
        event_service = EventService()
        participant_service = ParticipantService()

        # Any existing participant this import actually references (attendance or list
        # rows) must be eligible on protocol_date, or they'd silently vanish from this
        # protocol's roster despite being explicitly named in the imported document -
        # widen (never narrow) their joined_at/left_at window to include it. A None side
        # already means "unbounded" there, so it's left untouched.
        participant_cache: dict[int, Participant] = {}

        def _get_cached_participant(participant_id: int) -> Participant | None:
            if participant_id not in participant_cache:
                participant_cache[participant_id] = db.get(Participant, participant_id)
            return participant_cache[participant_id]

        def _widen_participant_window(participant_id: int, target_date: date) -> None:
            participant = _get_cached_participant(participant_id)
            if participant is None:
                return
            if participant.joined_at is not None and target_date < participant.joined_at:
                participant.joined_at = target_date
                db.add(participant)
            if participant.left_at is not None and target_date > participant.left_at:
                participant.left_at = target_date
                db.add(participant)

        # Append-only quality log (see WordImportSuggestionOutcome) - one row per
        # resolved decision where analyze() actually had a confident top suggestion to
        # compare against the human's final choice. Flushed once at the very end of
        # commit() alongside everything else (see db.add_all(outcome_rows) below).
        outcome_rows: list[WordImportSuggestionOutcome] = []
        # Negative-feedback half of the same original-vs-final comparison (see A.4 /
        # WordImportProfile.mapping_config_json["rejected_candidates"] docstring further
        # below) - keyed "{signal_prefix}:{normalized_context}" so analyze()'s read side
        # can demote a specific candidate the NEXT time this same context recurs, instead
        # of only ever reinforcing accepted matches like the other profile keys do.
        rejected_candidate_updates: dict[str, dict] = {}

        def _log_outcome(
            signal_type: str, suggested_score: float | None, original_id, final_id, rejection_key: str | None = None
        ) -> None:
            if original_id is None:
                # Nothing to compare against (e.g. analyze() found no candidate at all,
                # status "new") - not a "wrong suggestion", so not a logged outcome.
                return
            accepted = original_id == final_id
            outcome_rows.append(
                WordImportSuggestionOutcome(
                    tenant_id=tenant_id,
                    template_id=payload.template_id,
                    signal_type=signal_type,
                    # 0.0 sentinel for signal types that don't carry a real per-decision
                    # score yet (currently: name resolutions inside matrix/list/form
                    # blocks, see WordImportNameResolution) - still records the far more
                    # valuable accept/reject boolean without blocking on that gap.
                    suggested_score=suggested_score if suggested_score is not None else 0.0,
                    was_accepted=accepted,
                )
            )
            if not accepted and rejection_key is not None:
                entry = rejected_candidate_updates.setdefault(rejection_key, {"rejected": [], "chosen": None})
                if original_id not in entry["rejected"]:
                    entry["rejected"].append(original_id)
                entry["chosen"] = final_id

        template = db.get(Template, payload.template_id)
        cycle_assignments: list[CycleAssignment] | None = None
        if template is not None and template.cycle_config_id is not None:
            cycle_cfg = db.get(CycleConfig, template.cycle_config_id)
            if cycle_cfg is not None:
                cycle_year = get_cycle_year(payload.protocol_date, cycle_cfg.reset_month, cycle_cfg.reset_day)
                cycle_assignments = [CycleAssignment(cycle_config_id=cycle_cfg.id, cycle_year=cycle_year)]

        protocol_id = protocol_service.create_from_template(
            db,
            ProtocolCreateFromTemplate(
                template_id=payload.template_id,
                protocol_date=payload.protocol_date,
                event_id=None,
            ),
            tenant_id=tenant_id,
            created_by=user_id,
        )

        rows = list(
            db.execute(
                select(ProtocolElementBlock, ProtocolElement.template_element_id, ProtocolElement.id)
                .join(ProtocolElement, ProtocolElement.id == ProtocolElementBlock.protocol_element_id)
                .where(ProtocolElement.protocol_id == protocol_id)
            ).all()
        )
        block_by_key: dict[tuple[int, int], ProtocolElementBlock] = {}
        # (template_element_id, event_id) -> the event-repeat block instance already
        # auto-generated for that Event by create_from_template's own repeat logic
        # above (only events currently within its window get one there) - reused
        # instead of creating a duplicate block for the same Event.
        event_block_by_key: dict[tuple[int, int], ProtocolElementBlock] = {}
        protocol_element_id_by_template_element_id: dict[int, int] = {}
        attendance_blocks: list[ProtocolElementBlock] = []
        for block, template_element_id, protocol_element_id in rows:
            config = block.configuration_snapshot_json or {}
            source_sort_index = config.get("source_sort_index")
            if template_element_id is not None and source_sort_index is not None:
                block_by_key[(template_element_id, source_sort_index)] = block
            if template_element_id is not None:
                protocol_element_id_by_template_element_id[template_element_id] = protocol_element_id
                if config.get("repeat_source_type") == "event" and config.get("repeat_source_id") is not None:
                    event_block_by_key[(template_element_id, int(config["repeat_source_id"]))] = block
            if config and "attendance_entries" in config:
                attendance_blocks.append(block)

        def _get_or_create_event_repeat_block(template_element_id: int, event_id: int) -> ProtocolElementBlock | None:
            key = (template_element_id, event_id)
            block = event_block_by_key.get(key)
            if block is not None:
                return block
            protocol_element_id = protocol_element_id_by_template_element_id.get(template_element_id)
            if protocol_element_id is None:
                return None
            try:
                block = protocol_service.add_event_block_to_element(
                    db, protocol_element_id=protocol_element_id, event_id=event_id
                )
            except ValueError:
                return None
            event_block_by_key[key] = block
            return block

        # Same shape/purpose as attendance_name_updates/list_name_updates below - learned
        # from form-block participant(s) rows (e.g. "Organisation"/"Wer geht" on a
        # Scharanlässe-style block) so a repeated import of similarly-worded old
        # protocols reuses the same name resolution without re-asking.
        form_name_updates: dict[str, int] = {}
        # normalized raw_name -> newly created Participant.id, deduplicates a name typed
        # as "create new" appearing in more than one form-field row within this same
        # commit (e.g. the same person listed under both "Organisation" and "Wer geht").
        created_form_participants: dict[str, int] = {}
        for text_commit in payload.texts:
            if text_commit.template_element_id is None:
                continue
            if text_commit.is_event_repeat:
                if text_commit.linked_event_id is None:
                    # No Anlass chosen for this Rückblick-style section - never falls
                    # back to block_sort_index, which would silently write into an
                    # arbitrary other Event's block (see event_block_by_key above).
                    continue
                block = _get_or_create_event_repeat_block(text_commit.template_element_id, text_commit.linked_event_id)
            elif text_commit.block_sort_index is not None:
                block = block_by_key.get((text_commit.template_element_id, text_commit.block_sort_index))
            else:
                continue
            if block is None:
                continue
            if text_commit.is_form_block:
                config = dict(block.configuration_snapshot_json or {})
                fields_by_row_id = {field.row_id: field for field in text_commit.form_fields}
                updated_rows = []
                for row in config.get("rows") or []:
                    row_id = str(row.get("id"))
                    field = fields_by_row_id.get(row_id)
                    if field is None:
                        updated_rows.append(row)
                        continue
                    resolved_names = []
                    for name in field.names:
                        participant_id = name.participant_id
                        if participant_id is None and name.create_new:
                            cache_key = _normalize(name.raw_name)
                            participant_id = created_form_participants.get(cache_key)
                            if participant_id is None:
                                new_participant = participant_service.create_participant(
                                    db,
                                    ParticipantCreate(
                                        display_name=name.raw_name.strip(),
                                        # Scoped to exactly this protocol's date, same as a
                                        # newly created attendance participant above.
                                        joined_at=payload.protocol_date,
                                        left_at=payload.protocol_date,
                                    ),
                                    tenant_id=tenant_id,
                                )
                                db.flush()
                                participant_id = new_participant.id
                                created_form_participants[cache_key] = participant_id
                                db.add(TemplateParticipant(template_id=payload.template_id, participant_id=participant_id, exclude_from_attendance=False))
                        _log_outcome(
                            "name_match", name.originally_suggested_score, name.originally_suggested_participant_id, participant_id,
                            rejection_key=f"name:{_normalize(name.raw_name)}" if name.raw_name else None,
                        )
                        resolved_names.append(name.model_copy(update={"participant_id": participant_id}))
                        if participant_id is not None:
                            form_name_updates[_normalize(name.raw_name)] = participant_id
                            _widen_participant_window(participant_id, payload.protocol_date)
                    updated_rows.append({**row, **_resolved_value_json(field.row_type, field.raw_value, resolved_names)})
                config["rows"] = updated_rows
                block.configuration_snapshot_json = config
                db.add(block)
            else:
                protocol_text = db.execute(
                    select(ProtocolText).where(ProtocolText.protocol_element_block_id == block.id)
                ).scalar_one_or_none()
                if protocol_text is not None:
                    protocol_text.content = text_commit.content

        # Learned here (not straight from payload.attendance) so create_new rows contribute
        # their newly-created participant_id, and roster-only rows (raw_name == "", added in
        # analyze() for template participants missing from the document) don't pollute the
        # profile with a bogus ""-key override.
        attendance_name_updates: dict[str, int] = {}
        if payload.attendance and attendance_blocks:
            status_by_participant: dict[int, str] = {}
            created_participants: dict[int, str] = {}
            for entry in payload.attendance:
                participant_id = entry.participant_id
                if participant_id is None and entry.create_new:
                    new_participant = participant_service.create_participant(
                        db,
                        ParticipantCreate(
                            display_name=(entry.participant_name or entry.raw_name).strip(),
                            # Scoped to exactly this protocol's date, so the person shows up
                            # only here and not in other protocols' rosters until someone
                            # widens their membership window by hand.
                            joined_at=payload.protocol_date,
                            left_at=payload.protocol_date,
                        ),
                        tenant_id=tenant_id,
                    )
                    participant_id = new_participant.id
                    created_participants[participant_id] = new_participant.display_name
                    # Also add to the template's roster so future protocols/imports for this
                    # template already list them, not just this one-off protocol.
                    db.add(TemplateParticipant(template_id=payload.template_id, participant_id=participant_id, exclude_from_attendance=False))
                _log_outcome(
                    "participant_match", entry.originally_suggested_score, entry.originally_suggested_participant_id, participant_id,
                    rejection_key=f"name:{_normalize(entry.raw_name)}" if entry.raw_name else None,
                )
                if participant_id is None:
                    continue
                if participant_id not in created_participants:
                    _widen_participant_window(participant_id, payload.protocol_date)
                status_by_participant[participant_id] = entry.status
                if entry.raw_name:
                    attendance_name_updates[_normalize(entry.raw_name)] = participant_id
            if created_participants:
                db.flush()
            for attendance_block in attendance_blocks:
                old_entries = attendance_block.configuration_snapshot_json.get("attendance_entries", [])
                present_ids = {entry.get("participant_id") for entry in old_entries}
                # Build entirely new dicts/list rather than mutating old_entries (and its
                # entries) in place: those are the SAME objects SQLAlchemy's change-tracking
                # uses as its "before" snapshot for this JSON column, since JSONB columns
                # aren't Mutable-tracked here. Mutating them in place makes the reassigned
                # value compare equal to that snapshot at flush time, so the ORM silently
                # concludes nothing changed and never issues the UPDATE - status edits were
                # accepted in the wizard but never actually reached the database.
                new_entries = [
                    {**entry, "status": status_by_participant[entry["participant_id"]]}
                    if entry.get("participant_id") in status_by_participant
                    else entry
                    for entry in old_entries
                ]
                for missing_id, status in status_by_participant.items():
                    if missing_id in present_ids:
                        continue
                    # create_from_template built attendance_entries from the roster query
                    # BEFORE this method ran (and that query filters by participant_eligible_on)
                    # - so this participant is missing here either because they were just
                    # created above, or because their old joined_at/left_at window excluded
                    # protocol_date and only got widened to include it a few lines up. Either
                    # way: append instead of status-patch.
                    display_name = created_participants.get(missing_id)
                    if display_name is None:
                        existing_participant = _get_cached_participant(missing_id)
                        display_name = existing_participant.display_name if existing_participant else ""
                    new_entries.append({"participant_id": missing_id, "participant_name": display_name, "status": status})
                attendance_block.configuration_snapshot_json = {
                    **attendance_block.configuration_snapshot_json,
                    "attendance_entries": new_entries,
                }

        for event_commit in payload.events:
            if not event_commit.approved:
                continue
            _log_outcome(
                "event_match", event_commit.originally_suggested_score,
                event_commit.originally_suggested_event_id, event_commit.linked_event_id,
                rejection_key=f"event:{event_commit.final_date.isoformat()}|{_normalize(event_commit.final_title)}",
            )
            # tag/participant_count are only ever set for Matrix-sourced rows (see
            # WordImportEventCommit) - left out of the kwargs entirely for ordinary
            # Termine-table rows so neither is ever accidentally cleared/overwritten on
            # an existing Event's unrelated, independently-maintained value.
            extra_kwargs: dict = {}
            if event_commit.tag is not None:
                extra_kwargs["tag"] = event_commit.tag
            if event_commit.participant_count is not None:
                extra_kwargs["participant_count"] = event_commit.participant_count
            if event_commit.linked_event_id is None:
                event_service.create_event(
                    db,
                    EventCreate(
                        title=event_commit.final_title,
                        event_date=event_commit.final_date,
                        cycle_assignments=cycle_assignments,
                        **extra_kwargs,
                    ),
                    tenant_id=tenant_id,
                )
            else:
                event_service.update_event(
                    db,
                    event_commit.linked_event_id,
                    EventUpdate(
                        title=event_commit.final_title,
                        event_date=event_commit.final_date,
                        cycle_assignments=cycle_assignments,
                        **extra_kwargs,
                    ),
                )

        # Unlike lists, matrix cells are written straight into the live protocol's own
        # Matrix block's configuration_snapshot_json - there is no separate persisted
        # entity, and (see WordImportService plan) no freeze/refresh step touches matrix
        # blocks on the "abgeschlossen" transition below, so this can run immediately.
        matrix_name_updates: dict[str, int] = {}
        for matrix_commit in payload.matrices:
            if not matrix_commit.approved:
                continue
            _log_outcome(
                "matrix_column_match", matrix_commit.originally_suggested_score,
                matrix_commit.originally_suggested_column_key, matrix_commit.column_key,
                rejection_key=f"matrix_column:{matrix_commit.matrix_key}:{_normalize(matrix_commit.column_label)}",
            )
            try:
                template_element_id_str, sort_index_str = matrix_commit.matrix_key.split(":", 1)
                block_key = (int(template_element_id_str), int(sort_index_str))
            except (ValueError, AttributeError):
                continue
            block = block_by_key.get(block_key)
            if block is None:
                continue
            config = dict(block.configuration_snapshot_json or {})
            columns = [dict(column) for column in (config.get("columns") or [])]
            target_column = next((column for column in columns if str(column.get("id")) == matrix_commit.column_key), None)
            if target_column is None:
                target_column = {
                    "id": matrix_commit.column_key,
                    "title": matrix_commit.column_label,
                    "sort_index": len(columns),
                    "row_values": {},
                }
                columns.append(target_column)
            row_values = dict(target_column.get("row_values") or {})
            row_values[matrix_commit.row_id] = _resolved_value_json(matrix_commit.row_type, matrix_commit.raw_value, matrix_commit.names)
            target_column["row_values"] = row_values
            config["columns"] = columns
            block.configuration_snapshot_json = config
            db.add(block)
            for name_resolution in matrix_commit.names:
                _log_outcome(
                    "name_match", name_resolution.originally_suggested_score,
                    name_resolution.originally_suggested_participant_id, name_resolution.participant_id,
                    rejection_key=f"name:{_normalize(name_resolution.raw_name)}" if name_resolution.raw_name else None,
                )
                if name_resolution.participant_id is not None:
                    matrix_name_updates[_normalize(name_resolution.raw_name)] = name_resolution.participant_id
                    _widen_participant_window(name_resolution.participant_id, payload.protocol_date)

        # Lists are never written to the live ListEntry table (see WordImportListRowMapping
        # / WordImportListRowCommit docstrings) - collect the approved rows here and only
        # patch the protocol's own block snapshot(s) further below, *after* the status
        # transition to "abgeschlossen" has run its live-refreshing freeze step (otherwise
        # that freeze would immediately overwrite these historical values with today's data).
        list_name_updates: dict[str, int] = {}
        approved_list_commits: list[WordImportListRowCommit] = []
        definitions_cache: dict[int, ListDefinition] = {}
        for list_commit in payload.lists:
            if not list_commit.approved:
                continue
            definition = definitions_cache.get(list_commit.list_definition_id)
            if definition is None:
                definition = db.get(ListDefinition, list_commit.list_definition_id)
                if definition is None:
                    continue
                definitions_cache[list_commit.list_definition_id] = definition
            approved_list_commits.append(list_commit)
            _log_outcome(
                "list_entry_match", list_commit.originally_suggested_score,
                list_commit.originally_suggested_entry_id, list_commit.linked_entry_id,
                rejection_key=f"list_entry:{list_commit.list_definition_id}:{_normalize(list_commit.column_one_raw)}",
            )
            for name_resolution in list_commit.column_one_names + list_commit.column_two_names:
                _log_outcome(
                    "name_match", name_resolution.originally_suggested_score,
                    name_resolution.originally_suggested_participant_id, name_resolution.participant_id,
                    rejection_key=f"name:{_normalize(name_resolution.raw_name)}" if name_resolution.raw_name else None,
                )
                if name_resolution.participant_id is not None:
                    list_name_updates[_normalize(name_resolution.raw_name)] = name_resolution.participant_id
                    _widen_participant_window(name_resolution.participant_id, payload.protocol_date)

        heading_updates = {
            _normalize(tc.extracted_heading): {
                "template_element_id": tc.template_element_id,
                "block_sort_index": tc.block_sort_index,
            }
            for tc in payload.texts
            # Event-repeat headings (e.g. "Rückblick Elternabend") name a specific
            # Anlass, not a stable structural target - memoizing them would make a
            # differently-named Rückblick heading next time wrongly short-circuit
            # straight past _match_event_repeat_section in analyze().
            if tc.template_element_id is not None and tc.block_sort_index is not None and not tc.is_event_repeat
        }
        table_role_updates: dict[str, dict] = {}
        for tc in payload.tables:
            _log_outcome(
                "table_role", tc.originally_suggested_score, tc.originally_suggested_role, tc.role,
                rejection_key=f"table_role:{tc.header_signature}",
            )
            table_role_updates[tc.header_signature] = {
                "role": tc.role,
                "list_definition_id": tc.list_definition_id,
                "matrix_key": tc.matrix_key,
                "list_grouping_strategy": tc.list_grouping_strategy,
            }
        name_updates: dict[str, int] = {
            **attendance_name_updates,
            **list_name_updates,
            **form_name_updates,
            **matrix_name_updates,
        }
        if heading_updates or table_role_updates or name_updates or rejected_candidate_updates:
            profile = db.execute(
                select(WordImportProfile).where(
                    WordImportProfile.tenant_id == tenant_id, WordImportProfile.template_id == payload.template_id
                )
            ).scalar_one_or_none()
            if profile is None:
                profile = WordImportProfile(tenant_id=tenant_id, template_id=payload.template_id, mapping_config_json={})
                db.add(profile)
                db.flush()
            config = dict(profile.mapping_config_json or {})
            heading_map = dict(config.get("heading_to_target", {}))
            heading_map.update(heading_updates)
            table_map = dict(config.get("table_roles_by_signature", {}))
            table_map.update(table_role_updates)
            name_map = dict(config.get("participant_name_overrides", {}))
            name_map.update(name_updates)
            # Additive negative-feedback store (see A.4 plan / _log_outcome above) - one
            # entry per "{signal_prefix}:{normalized_context}" key, merging newly rejected
            # candidate ids into any list already recorded for that same context rather
            # than overwriting it, so a context that's been wrong more than once across
            # separate imports accumulates every distinct wrong candidate seen so far.
            rejected_map = dict(config.get("rejected_candidates", {}))
            for key, update in rejected_candidate_updates.items():
                existing_entry = dict(rejected_map.get(key) or {"rejected": [], "chosen": None})
                merged_rejected = list(existing_entry.get("rejected") or [])
                for candidate_id in update["rejected"]:
                    if candidate_id not in merged_rejected:
                        merged_rejected.append(candidate_id)
                rejected_map[key] = {"rejected": merged_rejected, "chosen": update["chosen"]}
            profile.mapping_config_json = {
                "heading_to_target": heading_map,
                "table_roles_by_signature": table_map,
                "participant_name_overrides": name_map,
                "rejected_candidates": rejected_map,
            }

        protocol_service.update_protocol(db, protocol_id, ProtocolUpdate(status="abgeschlossen"))

        if approved_list_commits:
            blocks_by_list_id: dict[int, list[ProtocolElementBlock]] = {}
            for block in db.execute(
                select(ProtocolElementBlock)
                .join(ProtocolElement, ProtocolElement.id == ProtocolElementBlock.protocol_element_id)
                .where(ProtocolElement.protocol_id == protocol_id)
            ).scalars():
                linked_list_id = (block.configuration_snapshot_json or {}).get("linked_list_id")
                if linked_list_id:
                    blocks_by_list_id.setdefault(int(linked_list_id), []).append(block)

            synthetic_id = 0
            for list_commit in approved_list_commits:
                target_blocks = blocks_by_list_id.get(list_commit.list_definition_id)
                if not target_blocks:
                    # Chosen template has no block linked to this list - no snapshot slot
                    # exists in this protocol, so nothing is written (never falls back to
                    # live, see WordImportListRowMapping.has_snapshot_target).
                    continue
                definition = definitions_cache[list_commit.list_definition_id]
                col1_value = _resolved_value_json(definition.column_one_value_type, list_commit.column_one_raw, list_commit.column_one_names)
                col2_value = _resolved_value_json(definition.column_two_value_type, list_commit.column_two_raw, list_commit.column_two_names)
                for block in target_blocks:
                    config = dict(block.configuration_snapshot_json or {})
                    list_snapshot = dict(config.get("list_snapshot") or {})
                    entries = list(list_snapshot.get("entries") or [])
                    target_index = (
                        next(
                            (i for i, entry in enumerate(entries) if isinstance(entry, dict) and entry.get("id") == list_commit.linked_entry_id),
                            None,
                        )
                        if list_commit.linked_entry_id is not None
                        else None
                    )
                    if target_index is not None:
                        entries[target_index] = {**entries[target_index], "column_one_value": col1_value, "column_two_value": col2_value}
                    else:
                        # No live counterpart (or it wasn't found in this block's frozen
                        # entries) - append a snapshot-only row with a synthetic negative id,
                        # never a real ListEntry id, so it can never collide with one.
                        synthetic_id -= 1
                        entries.append(
                            {
                                "id": list_commit.linked_entry_id if list_commit.linked_entry_id is not None else synthetic_id,
                                "sort_index": len(entries),
                                "column_one_value": col1_value,
                                "column_two_value": col2_value,
                            }
                        )
                    list_snapshot["entries"] = entries
                    config["list_snapshot"] = list_snapshot
                    block.configuration_snapshot_json = config
                    db.add(block)

        if outcome_rows:
            db.add_all(outcome_rows)

        db.commit()
        return protocol_id
